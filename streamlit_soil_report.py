import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from scipy import stats
import io
from pathlib import Path
import re
from datetime import datetime
import base64
import geopandas as gpd
from shapely.geometry import mapping, Point, Polygon
from shapely import wkb, wkt
from rasterio.features import geometry_mask
from affine import Affine
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from pykrige.ok import OrdinaryKriging

# PDF generation imports
import reportlab
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import tempfile
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv
import binascii
import sys

# Excel formatting imports
from openpyxl.styles import PatternFill

# Translations dictionary
TRANSLATIONS = {
    'pt': {
        'page_title': 'Relatório de Classificação de Solo',
        'main_title': '🌱 Relatório de Classificação de Solo',
        'subtitle': '### Sistema de Análise e Classificação Automatizada',
        'language': 'Idioma',
        'settings': '⚙️ Configurações',
        'upload_file': '📁 Upload do arquivo de dados',
        'upload_help': 'Faça upload do arquivo Excel ou CSV com os dados de solo. Arquivos em inglês serão automaticamente traduzidos!',
        'project_name': '📝 Nome do Projeto',
        'project_help': 'Nome que aparecerá no relatório',
        'default_project': 'Análise de Solo',
        'file_loaded': '✅ Arquivo carregado com sucesso!',
        'lines_found': 'linhas encontradas.',
        'data_preview': '👀 Prévia dos Dados',
        'columns_available': '**Colunas disponíveis:**',
        'column_mapping': '🔧 Mapeamento de Colunas',
        'parameter_column': 'Coluna de Parâmetros',
        'parameter_help': 'Selecione a coluna que contém os nomes dos parâmetros',
        'value_column': 'Coluna de Valores',
        'value_help': 'Selecione a coluna que contém os valores numéricos',
        'custom_parameters': '🎯 Parâmetros Customizados',
        'add_custom': 'Adicione parâmetros personalizados:',
        'parameter_name': 'Nome do Parâmetro',
        'parameter_type': 'Tipo',
        'add_parameter': '➕ Adicionar Parâmetro',
        'parameter_added': 'Parâmetro \'{}\' adicionado!',
        'run_classification': '🚀 Executar Classificação',
        'processing': 'Processando classificação...',
        'classification_completed': '✅ Classificação concluída!',
        'total_samples': '📊 Total de Amostras',
        'predominant_class': '🏆 Classificação Predominante',
        'unique_parameters': '🧪 Parâmetros Únicos',
        'visualizations': '📈 Visualizações',
        'classification_summary': 'Resumo das Classificações',
        'samples': 'Número de Amostras',
        'advanced_analysis': '🔬 Análise Estatística Avançada',
        'select_parameter': 'Selecione um parâmetro para análise:',
        'group_by': 'Agrupar por:',
        'group_help': 'Escolha uma coluna para agrupar a análise',
        'distribution': '📊 Distribuição',
        'box_plot': '📈 Box Plot',
        'statistics': '📋 Estatísticas',
        'comparison': '⚖️ Comparação',
        'box_plot_unavailable': 'Box plot não disponível para esta combinação',
        'statistical_summary': '📊 Resumo Estatístico',
        'select_grouping': 'Selecione colunas para agrupamento:',
        'grouping_help': 'Escolha uma ou mais colunas para calcular estatísticas agrupadas',
        'download_stats': '📥 Download Estatísticas (CSV)',
        'no_stats_available': 'Nenhuma estatística disponível para este parâmetro com o agrupamento selecionado',
        'select_one_column': 'Selecione pelo menos uma coluna para agrupamento',
        'comparison_between_groups': '⚖️ Comparação entre Grupos',
        'compare_by': 'Comparar por:',
        'statistic': 'Estatística:',
        'mean': 'Média',
        'std': 'Desvio Padrão',
        'median': 'Mediana',
        'numerical_values': 'Valores Numéricos:',
        'classification_breakdown': '📋 Detalhamento das Classificações',
        'statistical_overview': '📊 Visão Geral Estatística',
        'complete_statistics': '🔍 Estatísticas Completas por Grupo',
        'select_stat_groups': 'Selecione colunas para agrupamento estatístico:',
        'stat_groups_help': 'Estatísticas serão calculadas para cada combinação das colunas selecionadas',
        'complete_stats_table': '📈 Tabela de Estatísticas Completa',
        'download_complete_stats': '📥 Download Estatísticas Completas (CSV)',
        'parameter_summary': '🎯 Resumo por Parâmetro',
        'no_grouping_columns': 'Nenhuma coluna disponível para agrupamento. Certifique-se de que seu arquivo contém colunas categóricas como \'Tratamento\', \'Tempo\', etc.',
        'classified_data': '📋 Dados Classificados',
        'pdf_report_generation': '📄 Geração de Relatório PDF',
        'generate_pdf': '📄 Gerar Relatório PDF',
        'generating_pdf': 'Gerando relatório PDF...',
        'pdf_generated': '✅ Relatório PDF gerado com sucesso!',
        'download_pdf': '📥 Download do Relatório PDF',
        'pdf_error': '❌ Erro ao gerar PDF:',
        'generate_docx': '📝 Gerar Relatório DOCX',
        'generating_docx': 'Gerando relatório DOCX...',
        'docx_generated': '✅ Relatório DOCX gerado com sucesso!',
        'download_docx': '📥 Download do Relatório DOCX',
        'docx_error': '❌ Erro ao gerar DOCX:',
        'download_excel': '📊 Download Excel',
        'upload_instructions': '👆 Faça upload de um arquivo na barra lateral para começar',
        'instructions': '📋 Instruções de Uso:',
        'step1': '**Upload do Arquivo**: Carregue um arquivo Excel (.xlsx) ou CSV com seus dados de solo',
        'step2': '**Mapeamento**: Selecione as colunas corretas para parâmetros e valores',
        'step3': '**Configuração**: (Opcional) Adicione parâmetros customizados',
        'step4': '**Classificação**: Execute a classificação automática',
        'step5': '**Análise**: Visualize os resultados em gráficos e tabelas',
        'step6': '**Relatório**: Gere um relatório PDF profissional',
        'expected_format': '📊 Formato do Arquivo Esperado:',
        'format_description': 'O arquivo deve conter pelo menos duas colunas:',
        'parameter_col_desc': '**Coluna de Parâmetros**: Nome dos parâmetros (ex: "pH em CaCl2", "Matéria Orgânica")',
        'value_col_desc': '**Coluna de Valores**: Valores numéricos correspondentes',
        'supported_classifications': '🎯 Classificações Suportadas:',
        'mb_mbom': '**MB → MBom**: Muito Baixo → Muito Bom (para macronutrientes)',
        'b_malto': '**B → MAlto**: Baixo → Muito Alto (para micronutrientes/pH)',
        # Classification levels
        'Muito Baixo': 'Muito Baixo',
        'Baixo': 'Baixo',
        'Médio': 'Médio',
        'Bom': 'Bom',
        'Muito Bom': 'Muito Bom',
        'Alto': 'Alto',
        'Muito Alto': 'Muito Alto',
        'Classificação não definida': 'Classificação não definida',
        'Valor inválido': 'Valor inválido'
    },
    'en': {
        'page_title': 'Soil Classification Report',
        'main_title': '🌱 Soil Classification Report',
        'subtitle': '### Automated Analysis and Classification System',
        'language': 'Language',
        'settings': '⚙️ Settings',
        'upload_file': '📁 Upload data file',
        'upload_help': 'Upload Excel or CSV file with soil data. English files will be automatically translated!',
        'project_name': '📝 Project Name',
        'project_help': 'Name that will appear in the report',
        'default_project': 'Soil Analysis',
        'file_loaded': '✅ File loaded successfully!',
        'lines_found': 'lines found.',
        'data_preview': '👀 Data Preview',
        'columns_available': '**Available columns:**',
        'column_mapping': '🔧 Column Mapping',
        'parameter_column': 'Parameter Column',
        'parameter_help': 'Select the column containing parameter names',
        'value_column': 'Value Column',
        'value_help': 'Select the column containing numerical values',
        'custom_parameters': '🎯 Custom Parameters',
        'add_custom': 'Add custom parameters:',
        'parameter_name': 'Parameter Name',
        'parameter_type': 'Type',
        'add_parameter': '➕ Add Parameter',
        'parameter_added': 'Parameter \'{}\' added!',
        'run_classification': '🚀 Run Classification',
        'processing': 'Processing classification...',
        'classification_completed': '✅ Classification completed!',
        'total_samples': '📊 Total Samples',
        'predominant_class': '🏆 Predominant Classification',
        'unique_parameters': '🧪 Unique Parameters',
        'visualizations': '📈 Visualizations',
        'classification_summary': 'Classification Summary',
        'samples': 'Number of Samples',
        'advanced_analysis': '🔬 Advanced Statistical Analysis',
        'select_parameter': 'Select a parameter for analysis:',
        'group_by': 'Group by:',
        'group_help': 'Choose a column to group the analysis',
        'distribution': '📊 Distribution',
        'box_plot': '📈 Box Plot',
        'statistics': '📋 Statistics',
        'comparison': '⚖️ Comparison',
        'box_plot_unavailable': 'Box plot not available for this combination',
        'statistical_summary': '📊 Statistical Summary',
        'select_grouping': 'Select columns for grouping:',
        'grouping_help': 'Choose one or more columns to calculate grouped statistics',
        'download_stats': '📥 Download Statistics (CSV)',
        'no_stats_available': 'No statistics available for this parameter with selected grouping',
        'select_one_column': 'Select at least one column for grouping',
        'comparison_between_groups': '⚖️ Comparison between Groups',
        'compare_by': 'Compare by:',
        'statistic': 'Statistic:',
        'mean': 'Mean',
        'std': 'Standard Deviation',
        'median': 'Median',
        'numerical_values': 'Numerical Values:',
        'classification_breakdown': '📋 Classification Breakdown',
        'statistical_overview': '📊 Statistical Overview',
        'complete_statistics': '🔍 Complete Statistics by Group',
        'select_stat_groups': 'Select columns for statistical grouping:',
        'stat_groups_help': 'Statistics will be calculated for each combination of selected columns',
        'complete_stats_table': '📈 Complete Statistics Table',
        'download_complete_stats': '📥 Download Complete Statistics (CSV)',
        'parameter_summary': '🎯 Summary by Parameter',
        'no_grouping_columns': 'No columns available for grouping. Make sure your file contains categorical columns like \'Treatment\', \'Time\', etc.',
        'classified_data': '📋 Classified Data',
        'pdf_report_generation': '📄 PDF Report Generation',
        'generate_pdf': '📄 Generate PDF Report',
        'generating_pdf': 'Generating PDF report...',
        'pdf_generated': '✅ PDF report generated successfully!',
        'download_pdf': '📥 Download PDF Report',
        'pdf_error': '❌ Error generating PDF:',
        'generate_docx': '📝 Generate DOCX Report',
        'generating_docx': 'Generating DOCX report...',
        'docx_generated': '✅ DOCX report generated successfully!',
        'download_docx': '📥 Download DOCX Report',
        'docx_error': '❌ Error generating DOCX:',
        'download_excel': '📊 Download Excel',
        'upload_instructions': '👆 Upload a file in the sidebar to start',
        'instructions': '📋 Usage Instructions:',
        'step1': '**File Upload**: Upload an Excel (.xlsx) or CSV file with your soil data',
        'step2': '**Mapping**: Select the correct columns for parameters and values',
        'step3': '**Configuration**: (Optional) Add custom parameters',
        'step4': '**Classification**: Run automatic classification',
        'step5': '**Analysis**: View results in charts and tables',
        'step6': '**Report**: Generate a professional PDF report',
        'expected_format': '📊 Expected File Format:',
        'format_description': 'The file should contain at least two columns:',
        'parameter_col_desc': '**Parameter Column**: Parameter names (e.g., "pH in CaCl2", "Organic Matter")',
        'value_col_desc': '**Value Column**: Corresponding numerical values',
        'supported_classifications': '🎯 Supported Classifications:',
        'mb_mbom': '**MB → MBom**: Very Low → Very Good (for macronutrients)',
        'b_malto': '**B → MAlto**: Low → Very High (for micronutrients/pH)',
        # Classification levels
        'Muito Baixo': 'Very Low',
        'Baixo': 'Low',
        'Médio': 'Medium',
        'Bom': 'Good',
        'Muito Bom': 'Very Good',
        'Alto': 'High',
        'Muito Alto': 'Very High',
        'Classificação não definida': 'Classification not defined',
        'Valor inválido': 'Invalid value'
    }
}

# Language mapping for statistics
STAT_NAMES = {
    'pt': {'mean': 'Média', 'std': 'Desvio Padrão', 'median': 'Mediana'},
    'en': {'mean': 'Mean', 'std': 'Standard Deviation', 'median': 'Median'    }
}

# Database retrieval functions - using agbenefits pipeline connection method
def setup_database_connection():
    """Setup database connection using the same method as agbenefits pipeline"""
    try:
        # Load environment variables first
        load_dotenv()
        
        # Import the local database utilities
        from utils.db import get_terradot_db_session, read_pd_from_db_sql
        
        # Test the connection
        connection = get_terradot_db_session()
        if connection is None:
            st.error("❌ Database connection failed - check environment variables")
            st.info("Required environment variables: DB_HOST, DB_NAME, DB_USER")
            return None
            
        connection.close()
        
        st.success("✅ Database connection configured using agbenefits pipeline method")
        return True
        
    except ImportError as e:
        st.error(f"❌ Could not import database utilities: {str(e)}")
        st.info("Make sure the utils/db.py file is in the correct location")
        return None
    except Exception as e:
        st.error(f"❌ Database connection failed: {str(e)}")
        return None

def execute_sql_query(query, connection):
    """Execute SQL query and return DataFrame"""
    try:
        return pd.read_sql(query, connection)
    except Exception as e:
        st.error(f"Error executing query: {str(e)}")
        return pd.DataFrame()

def retrieve_soil_samples_from_db(field_id):
    """Retrieve soil samples from database using agbenefits pipeline method"""
    try:
        # Import the local database utilities
        from utils.db import get_terradot_db_session, read_pd_from_db_sql
        
        # Read SQL queries
        composite_query_path = "agbenefits_get_composite_samples.sql"
        noncomposite_query_path = "agbenefits_get_NoNcomposite_samples.sql"
        
        # Load composite samples query
        with open(composite_query_path, 'r') as file:
            composite_query = file.read().format(field_id=field_id)
        
        # Load non-composite samples query
        with open(noncomposite_query_path, 'r') as file:
            noncomposite_query = file.read().format(field_id=field_id)
        
        # Execute queries using agbenefits pipeline method
        connection = get_terradot_db_session()
        try:
            # Try composite samples first
            composite_df = read_pd_from_db_sql(composite_query, connection)
            st.info(f"Number of composite samples: {len(composite_df)}")
            
            # Try non-composite samples
            noncomposite_df = read_pd_from_db_sql(noncomposite_query, connection)
            st.info(f"Number of non-composite samples: {len(noncomposite_df)}")
            
            # Combine results (same logic as agbenefits pipeline)
            if len(composite_df) != 0 and len(noncomposite_df) != 0:
                combined_df = pd.concat([composite_df, noncomposite_df], ignore_index=True)
            elif len(composite_df) != 0:
                combined_df = composite_df
            elif len(noncomposite_df) != 0:
                combined_df = noncomposite_df
            else:
                st.warning(f"No soil samples found for field_id: {field_id}")
                combined_df = pd.DataFrame()
                
        finally:
            connection.close()
            
        return combined_df
        
    except Exception as e:
        st.error(f"Error retrieving soil samples: {str(e)}")
        return pd.DataFrame()

def retrieve_field_boundaries_from_db(field_id):
    """Retrieve field boundaries from database using agbenefits pipeline method"""
    try:
        # Import the local database utilities
        from utils.db import get_terradot_db_session, read_pd_from_db_sql
        
        # Read plot boundaries SQL query (same as agbenefits pipeline)
        boundary_query_path = "get_plot_boundaries.sql"
        with open(boundary_query_path, 'r') as file:
            boundary_query = file.read().format(field_id=field_id)
        
        # Execute query using agbenefits pipeline method
        connection = get_terradot_db_session()
        try:
            boundary_df = read_pd_from_db_sql(boundary_query, connection)
            
            if boundary_df.empty:
                return None
            
            # Convert WKB to geometry
            geometries = []
            for wkb_data in boundary_df['boundary']:
                try:
                    if isinstance(wkb_data, str) and wkb_data.startswith('\\x'):
                        wkb_bytes = binascii.unhexlify(wkb_data[2:])
                    else:
                        wkb_bytes = binascii.unhexlify(wkb_data)
                    geometries.append(wkb.loads(wkb_bytes))
                except:
                    geometries.append(None)
            
            # Create GeoDataFrame
            gdf = gpd.GeoDataFrame(boundary_df, geometry=geometries, crs='EPSG:4326')
            return gdf
            
        finally:
            connection.close()
        
    except Exception as e:
        st.error(f"Error retrieving field boundaries: {str(e)}")
        return None

# Set page config
st.set_page_config(
    page_title="Soil Classification Report",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize language in session state
if 'language' not in st.session_state:
    st.session_state.language = 'en'

def t(key):
    """Translation function"""
    return TRANSLATIONS[st.session_state.language].get(key, key)

def translate_classification(classification):
    """Translate classification levels"""
    return TRANSLATIONS[st.session_state.language].get(classification, classification)

def translate_parameter_for_display(param_name):
    """Translate parameter name for display purposes based on current language"""
    if st.session_state.language == 'pt':
        # If language is Portuguese and parameter is in English, translate it
        parameter_translations = get_parameter_translations()
        if param_name.lower() in parameter_translations:
            return parameter_translations[param_name.lower()]
        return param_name
    else:
        # If language is English and parameter is in Portuguese, translate back to English
        parameter_translations = get_parameter_translations()
        # Create reverse mapping
        reverse_translations = {v: k for k, v in parameter_translations.items()}
        if param_name in reverse_translations:
            return reverse_translations[param_name].title()  # Title case for display
        return param_name

def translate_column_for_display(col_name):
    """Translate column name for display purposes based on current language"""
    if st.session_state.language == 'pt':
        # If language is Portuguese and column is in English, translate it
        column_translations = get_column_translations()
        if col_name in column_translations:
            return column_translations[col_name]
        return col_name
    else:
        # If language is English and column is in Portuguese, translate back to English
        column_translations = get_column_translations()
        # Create reverse mapping
        reverse_translations = {v: k for k, v in column_translations.items()}
        if col_name in reverse_translations:
            return reverse_translations[col_name].replace('_', ' ').title()  # Title case for display
        return col_name

def get_grouping_columns_with_display_names(df, excluded_cols):
    """Get available grouping columns with their display names"""
    column_translations = get_column_translations()
    reverse_translations = {v: k for k, v in column_translations.items()}
    
    # Get available columns for grouping
    available_cols = [col for col in df.columns 
                     if col not in excluded_cols 
                     and df[col].dtype == 'object']
    
    # Create mapping of display names to actual column names
    display_to_actual = {}
    for col in available_cols:
        display_name = translate_column_for_display(col)
        display_to_actual[display_name] = col
    
    return display_to_actual

# Custom CSS
st.markdown("""
<style>
    .main > div {
        padding-top: 2rem;
    }
    .stAlert {
        margin-top: 1rem;
    }
    .metric-container {
        background-color: #f0f2f6;
        border-radius: 10px;
        padding: 10px;
        margin: 5px;
    }
    .kriging-map-container {
        max-height: 500px;
        overflow: hidden;
    }
    .stSlider > div > div > div {
        max-width: 400px;
    }
</style>
""", unsafe_allow_html=True)

# Configuration constants (keep original Portuguese classifications for data processing)
PARAMS_MB_DEFAULT = {
    # Matéria Orgânica
    "Matéria orgânica (MO)": [(0, 0.7, "Muito Baixo"), (0.8, 1.5, "Baixo"), (1.6, 2.0, "Médio"), (2.1, 3.0, "Bom"), (3.1, float('inf'), "Muito Bom")],
    "Carbono orgânico total": [(0, 8, "Muito Baixo"), (8, 15, "Baixo"), (15, 25, "Médio"), (25, 35, "Bom"), (35, float('inf'), "Muito Bom")],
    
    # Macronutrientes
    "Cálcio trocável (Ca2+)": [(0, 0.4, "Muito Baixo"), (0.4, 1.20, "Baixo"), (1.21, 2.40, "Médio"), (2.41, 4.00, "Bom"), (4.01, float('inf'), "Muito Bom")],
    "Magnésio trocável (Mg2+)": [(0, 0.15, "Muito Baixo"), (0.16, 0.45, "Baixo"), (0.46, 0.90, "Médio"), (0.91, 1.50, "Bom"), (1.51, float('inf'), "Muito Bom")],
    "Potássio trocável (K+)": [(0, 25.0, "Muito Baixo"), (26.0, 50.0, "Baixo"), (51.0, 80.0, "Médio"), (81.0, 100.0, "Bom"), (100.1, float('inf'), "Muito Bom")],
    
    # Fósforo (Resina, Sequeiro)
    "P - disponível": [(0, 5.0, "Muito Baixo"), (5.0, 8.0, "Baixo"), (8.0, 14.0, "Médio"), (14.0, 20.0, "Bom"), (20.0, float('inf'), "Muito Bom")],
    "P disponível (Resina, Sequeiro)": [(0, 5.0, "Muito Baixo"), (5.0, 8.0, "Baixo"), (8.0, 14.0, "Médio"), (14.0, 20.0, "Bom"), (20.0, float('inf'), "Muito Bom")],
    "P disponível (Resina, Irrigado)": [(0, 8.0, "Muito Baixo"), (8.0, 14.0, "Baixo"), (14.0, 20.0, "Médio"), (20.0, 35.0, "Bom"), (35.0, float('inf'), "Muito Bom")],
}

PARAMS_BA_DEFAULT = {
    # pH
    "pH em CaCl₂": [(0, 4.4, "Baixo"), (4.4, 4.8, "Médio"), (4.8, 5.5, "Bom"), (5.5, 5.8, "Alto"), (5.8, float('inf'), "Muito Alto")],
    "pH em H₂O": [(0, 5.1, "Baixo"), (5.1, 5.5, "Médio"), (5.5, 6.3, "Bom"), (6.3, 6.6, "Alto"), (6.6, float('inf'), "Muito Alto")],
    
    # Saturações
    "Saturação por bases (V%)": [(0, 20.0, "Baixo"), (20.0, 35.0, "Médio"), (35.0, 60.0, "Bom"), (60.0, 70.0, "Alto"), (70.0, float('inf'), "Muito Alto")],
    "Saturação por alumínio (m%)": [(60.1, float('inf'), "Muito Alto"), (20.1, 60.0, "Alto"), (0, 20.0, "Muito Baixo")],  # Invertido - menor é melhor
    
    # CTC
    "CTC efetiva (t)": [(0, 0.8, "Baixo"), (0.8, 2.30, "Médio"), (2.30, 4.60, "Bom"), (4.60, 8.00, "Alto"), (8.00, float('inf'), "Muito Alto")],
    "CTC a pH 7,0 (T)": [(0, 1.6, "Baixo"), (1.6, 4.30, "Médio"), (4.30, 8.60, "Bom"), (8.60, 15.00, "Alto"), (15.00, float('inf'), "Muito Alto")],
    
    # Acidez (invertido - menor é melhor)
    "Acidez trocável (Al3+)": [(2.01, float('inf'), "Muito Alto"), (1.01, 2.00, "Alto"), (0.51, 1.00, "Médio"), (0.21, 0.50, "Baixo"), (0, 0.20, "Muito Baixo")],
    "Acidez potencial (H+Al)": [(9.01, float('inf'), "Muito Alto"), (5.01, 9.00, "Alto"), (2.51, 5.00, "Médio"), (1.01, 2.50, "Baixo"), (0, 1.0, "Muito Baixo")],
    
    # Micronutrientes
    "Cobre (Cu)": [(0, 0.3, "Baixo"), (0.3, 0.7, "Médio"), (0.7, 1.2, "Bom"), (1.2, 1.8, "Alto"), (1.8, float('inf'), "Muito Alto")],
    "Ferro (Fe)": [(0, 8.0, "Baixo"), (8.0, 18.0, "Médio"), (18.0, 30.0, "Bom"), (30.0, 45.0, "Alto"), (45.0, float('inf'), "Muito Alto")],
    "Manganês (Mn)": [(0, 2.0, "Baixo"), (2.0, 5.0, "Médio"), (5.0, 8.0, "Bom"), (8.0, 12.0, "Alto"), (12.0, float('inf'), "Muito Alto")],
    "Zinco (Zn)": [(0, 0.4, "Baixo"), (0.4, 0.9, "Médio"), (0.9, 1.5, "Bom"), (1.5, 2.2, "Alto"), (2.2, float('inf'), "Muito Alto")],
    "Enxofre (S)": [(0, 2.0, "Baixo"), (2.0, 4.0, "Médio"), (4.0, 10.0, "Bom"), (10.0, 12.0, "Alto"), (12.0, float('inf'), "Muito Alto")],
    "Boro (B)": [(0, 0.2, "Baixo"), (0.16, 0.35, "Médio"), (0.36, 0.60, "Bom"), (0.61, 0.90, "Alto"), (0.90, float('inf'), "Muito Alto")],
}

class SoilClassifier:
    def __init__(self):
        self.params_mb = PARAMS_MB_DEFAULT.copy()
        self.params_ba = PARAMS_BA_DEFAULT.copy()
    
    def add_custom_parameter(self, param_name, ranges, param_type="MB"):
        """Add custom parameter classification ranges"""
        if param_type == "MB":
            self.params_mb[param_name] = ranges
        else:
            self.params_ba[param_name] = ranges
    
    def classify_value(self, param_name, value):
        """Classify a single value - handles both English and Portuguese parameter names"""
        try:
            value = float(value)
        except:
            return "Valor inválido"
        
        # Get parameter translations
        parameter_translations = get_parameter_translations()
        
        # If parameter is in English, translate it to Portuguese for classification
        portuguese_param_name = param_name
        if param_name.lower() in parameter_translations:
            portuguese_param_name = parameter_translations[param_name.lower()]
        
        # Check MB parameters first
        if portuguese_param_name in self.params_mb:
            for min_val, max_val, classification in self.params_mb[portuguese_param_name]:
                if min_val <= value < max_val:
                    return classification
            # If no range found, assign to the nearest category based on distance
            ranges = self.params_mb[portuguese_param_name]
            return self._find_nearest_classification(value, ranges)
        
        # Check BA parameters
        if portuguese_param_name in self.params_ba:
            for min_val, max_val, classification in self.params_ba[portuguese_param_name]:
                if min_val <= value < max_val:
                    return classification
            # If no range found, assign to the nearest category based on distance
            ranges = self.params_ba[portuguese_param_name]
            return self._find_nearest_classification(value, ranges)
        
        return "Classificação não definida"
    
    def _find_nearest_classification(self, value, ranges):
        """Find the classification with the minimum distance to the value"""
        min_distance = float('inf')
        nearest_classification = None
        
        for min_val, max_val, classification in ranges:
            # Calculate distance to range center
            range_center = (min_val + max_val) / 2
            distance = abs(value - range_center)
            
            if distance < min_distance:
                min_distance = distance
                nearest_classification = classification
        
        return nearest_classification if nearest_classification else ranges[-1][2]
    
    def classify_dataframe(self, df, param_col="Parâmetro", value_col="Resultado numérico"):
        """Classify entire dataframe"""
        if param_col not in df.columns or value_col not in df.columns:
            st.error(f"Colunas '{param_col}' ou '{value_col}' não encontradas no arquivo")
            return df
        
        df = df.copy()
        df["Classificação"] = df.apply(
            lambda row: self.classify_value(row[param_col], row[value_col]), 
            axis=1
        )
        return df

def create_classification_summary(df):
    """Create summary statistics for classifications"""
    if "Classificação" not in df.columns:
        return {}
    
    summary = df["Classificação"].value_counts().to_dict()
    total = len(df)
    
    return {
        "total_samples": total,
        "classification_counts": summary,
        "classification_percentages": {k: (v/total)*100 for k, v in summary.items()}
    }

def create_kde_curve(data, x_range, bandwidth=None):
    """Create KDE curve data"""
    if len(data) < 2:
        return x_range, np.zeros_like(x_range)
    
    try:
        kde = stats.gaussian_kde(data, bw_method=bandwidth)
        density = kde(x_range)
        return x_range, density
    except:
        # Fallback to simple histogram if KDE fails
        hist, bin_edges = np.histogram(data, bins=50, density=True)
        bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2
        return bin_centers, hist

def create_parameter_chart(df, param_name, group_by_cols=None, param_col="Parâmetro", value_col="Resultado numérico", separate_by_classification=False):
    """Create KDE density curves for specific parameter with optional grouping"""
    if param_col not in df.columns:
        return None
    
    param_data = df[df[param_col] == param_name].copy()
    if param_data.empty:
        return None
    
    # Translate classifications for display
    param_data["Classificação_Display"] = param_data["Classificação"].apply(translate_classification)
    
    # Translate parameter name for display
    display_param_name = translate_parameter_for_display(param_name)
    
    # Get overall data range for x-axis
    all_values = param_data[value_col].dropna()
    if all_values.empty:
        return None
    
    x_min, x_max = all_values.min(), all_values.max()
    x_range = np.linspace(x_min, x_max, 200)
    
    # Handle subplot creation if separating by classification
    if separate_by_classification:
        classifications = param_data["Classificação_Display"].unique()
        n_classifications = len(classifications)
        
        if n_classifications == 0:
            return None
        
        # Create subplots
        from plotly.subplots import make_subplots
        fig = make_subplots(
            rows=1, cols=n_classifications,
            subplot_titles=classifications,
            shared_yaxes=True,
            horizontal_spacing=0.05
        )
    else:
        # Create single figure
        fig = go.Figure()
    
    # Handle multiple grouping columns
    if group_by_cols:
        # Filter to only include columns that exist in the data
        valid_group_cols = [col for col in group_by_cols if col in param_data.columns]
        
        if valid_group_cols:
            # Create a combined grouping column for multiple filters
            if len(valid_group_cols) == 1:
                combined_group_col = valid_group_cols[0]
            else:
                # Create a combined grouping column
                param_data['Combined_Group'] = param_data[valid_group_cols].apply(
                    lambda row: ' | '.join([f"{col}: {row[col]}" for col in valid_group_cols]), 
                    axis=1
                )
                combined_group_col = 'Combined_Group'
            
            # Create plots for each classification
            classifications = param_data["Classificação_Display"].unique()
            
            for i, classification in enumerate(classifications):
                classification_data = param_data[param_data["Classificação_Display"] == classification]
                groups = classification_data[combined_group_col].unique()
                
                for group in groups:
                    group_data = classification_data[classification_data[combined_group_col] == group]
                    values = group_data[value_col].dropna()
                    
                    if len(values) > 1:
                        x_vals, density = create_kde_curve(values, x_range)
                        
                        if separate_by_classification:
                            # Add to specific subplot
                            fig.add_trace(go.Scatter(
                                x=x_vals,
                                y=density,
                                mode='lines',
                                fill='tozeroy',
                                name=f"{group}",
                                opacity=0.7,
                                line=dict(width=2),
                                showlegend=(i == 0)  # Only show legend for first subplot
                            ), row=1, col=i+1)
                        else:
                            # Add to single plot
                            fig.add_trace(go.Scatter(
                                x=x_vals,
                                y=density,
                                mode='lines',
                                fill='tozeroy',
                                name=f"{classification} - {group}",
                                opacity=0.7,
                                line=dict(width=2)
                            ))
            
            if separate_by_classification:
                fig.update_layout(
                title=f"{t('distribution')} - {display_param_name} por {' + '.join(valid_group_cols)}",
                    height=400
                )
                fig.update_xaxes(title_text="Valor")
                fig.update_yaxes(title_text="Density")
            else:
                fig.update_layout(
                    title=f"{t('distribution')} - {display_param_name} por {' + '.join(valid_group_cols)}",
                    xaxis_title="Valor",
                    yaxis_title="Density",
                    height=400
                )
        else:
            # Fallback to classification-only grouping
            classifications = param_data["Classificação_Display"].unique()
            
            for i, classification in enumerate(classifications):
                classification_data = param_data[param_data["Classificação_Display"] == classification]
                values = classification_data[value_col].dropna()
                
                if len(values) > 1:
                    x_vals, density = create_kde_curve(values, x_range)
                    
                    if separate_by_classification:
                        # Add to specific subplot
                        fig.add_trace(go.Scatter(
                            x=x_vals,
                            y=density,
                            mode='lines',
                            fill='tozeroy',
                            name=classification,
                            opacity=0.7,
                            line=dict(width=2),
                            showlegend=False  # No legend needed for single classification per subplot
                        ), row=1, col=i+1)
                else:
                        # Add to single plot
                        fig.add_trace(go.Scatter(
                            x=x_vals,
                            y=density,
                            mode='lines',
                            fill='tozeroy',
                            name=classification,
                            opacity=0.7,
                            line=dict(width=2)
                        ))
        
        if separate_by_classification:
            fig.update_layout(
                title=f"{t('distribution')} - {display_param_name}",
                height=400
            )
            fig.update_xaxes(title_text="Valor")
            fig.update_yaxes(title_text="Density")
        else:
            fig.update_layout(
                title=f"{t('distribution')} - {display_param_name}",
                xaxis_title="Valor",
                yaxis_title="Density",
                height=400
            )
    else:
        # No valid group columns, create default KDE plot by classification
        classifications = param_data["Classificação_Display"].unique()
        
        for i, classification in enumerate(classifications):
            classification_data = param_data[param_data["Classificação_Display"] == classification]
            values = classification_data[value_col].dropna()
            
            if len(values) > 1:
                    x_vals, density = create_kde_curve(values, x_range)
                    
                    if separate_by_classification:
                        # Add to specific subplot
                        fig.add_trace(go.Scatter(
                            x=x_vals,
                            y=density,
                            mode='lines',
                            fill='tozeroy',
                            name=classification,
                            opacity=0.7,
                            line=dict(width=2),
                            showlegend=False  # No legend needed for single classification per subplot
                        ), row=1, col=i+1)
                    else:
                        # Add to single plot
                        fig.add_trace(go.Scatter(
                                x=x_vals,
                                y=density,
                                mode='lines',
                                fill='tozeroy',
                                name=classification,
                                opacity=0.7,
                                line=dict(width=2)
                            ))
        
        if separate_by_classification:
            fig.update_layout(
                title=f"{t('distribution')} - {display_param_name}",
                height=400
            )
            fig.update_xaxes(title_text="Valor")
            fig.update_yaxes(title_text="Density")
        else:
            fig.update_layout(
                title=f"{t('distribution')} - {display_param_name}",
                xaxis_title="Valor",
                yaxis_title="Density",
                height=400
            )
    
    return fig

def create_spatial_plot(gdf, param_name, polygon_gdf=None, purpose_filter=None, depth_filter=None, 
                       param_col="Parâmetro", value_col="Resultado numérico", 
                       purpose_col="sampling_plan_purpose", depth_col="depth_range_bottom_m",
                       cmap="viridis", point_size=100):
    """Create spatial plot for a specific parameter with optional filters"""
    if param_col not in gdf.columns or value_col not in gdf.columns:
        return None
    
    # Filter data for the parameter
    param_data = gdf[gdf[param_col] == param_name].copy()
    if param_data.empty:
        return None
    
    # Apply filters
    if purpose_filter and purpose_filter != "All" and purpose_col in param_data.columns:
        param_data = param_data[param_data[purpose_col] == purpose_filter]
    
    if depth_filter and depth_filter != "All" and depth_col in param_data.columns:
        param_data = param_data[param_data[depth_col] == depth_filter]
    
    if param_data.empty:
        return None
    
    # Convert numeric values
    param_data[value_col] = pd.to_numeric(param_data[value_col], errors='coerce')
    param_data = param_data.dropna(subset=[value_col])
    
    if param_data.empty:
        return None
    
    # Create figure
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    
    # Get value range for consistent coloring
    vmin = param_data[value_col].min()
    vmax = param_data[value_col].max()
    
    # Handle case where all values are the same
    if vmin == vmax:
        vmin -= 0.01
        vmax += 0.01
    
    # Plot polygon boundaries if available
    if polygon_gdf is not None and not polygon_gdf.empty:
        # Ensure same CRS
        if polygon_gdf.crs != param_data.crs:
            polygon_gdf = polygon_gdf.to_crs(param_data.crs)
        
        # Plot boundaries
        polygon_gdf.boundary.plot(ax=ax, color="black", linewidth=2)
        
        # Add plot type labels if available
        if "plot_type" in polygon_gdf.columns:
            for _, row in polygon_gdf.iterrows():
                if row.geometry is not None and not row.geometry.is_empty:
                    centroid = row.geometry.centroid
                    label = str(row["plot_type"])[0].upper()
                    ax.text(
                        centroid.x, centroid.y, label,
                        ha="center", va="center",
                        fontsize=14, fontweight="bold", color="black",
                        bbox=dict(facecolor="white", alpha=0.7, boxstyle="circle,pad=0.3")
                    )
    
    # Plot points
    scatter = param_data.plot(
        ax=ax,
        column=value_col,
        cmap=cmap,
        markersize=point_size,
        alpha=0.8,
        vmin=vmin,
        vmax=vmax,
        legend=False,
        edgecolor="k",
        linewidth=0.5
    )
    
    # Set equal aspect ratio and labels
    ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel("Longitude", fontsize=12)
    ax.set_ylabel("Latitude", fontsize=12)
    
    # Title
    display_param_name = translate_parameter_for_display(param_name)
    title_parts = [display_param_name]
    
    if purpose_filter and purpose_filter != "All":
        title_parts.append(f"Purpose: {purpose_filter}")
    
    if depth_filter and depth_filter != "All":
        title_parts.append(f"Depth: {depth_filter}")
    
    title = " - ".join(title_parts)
    ax.set_title(title, fontsize=14, fontweight="bold")
    
    # Add colorbar
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
    sm._A = []
    cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
    cbar.set_label(value_col, fontsize=12)
    
    plt.tight_layout()
    return fig

def create_spatial_comparison_plot(gdf, param_name, polygon_gdf=None, purposes=("PRE_APPLICATION", "CREDIT_SAMPLING_1"),
                                 param_col="Parâmetro", value_col="Resultado numérico", 
                                 purpose_col="sampling_plan_purpose", depth_col="depth_range_bottom_m",
                                 cmap="viridis", point_size=100):
    """Create spatial comparison plot with two panels for different purposes"""
    if param_col not in gdf.columns or value_col not in gdf.columns:
        return None
    
    # Filter data for the parameter
    param_data = gdf[gdf[param_col] == param_name].copy()
    if param_data.empty:
        return None
    
    # Convert numeric values
    param_data[value_col] = pd.to_numeric(param_data[value_col], errors='coerce')
    param_data = param_data.dropna(subset=[value_col])
    
    if param_data.empty:
        return None
    
    # Get global value range for consistent coloring across panels
    vmin = param_data[value_col].min()
    vmax = param_data[value_col].max()
    
    if vmin == vmax:
        vmin -= 0.01
        vmax += 0.01
    
    # Create figure with two subplots
    fig, axes = plt.subplots(1, 2, figsize=(16, 8))
    
    for i, purpose in enumerate(purposes):
        ax = axes[i]
        
        # Filter data for this purpose
        purpose_data = param_data[param_data[purpose_col] == purpose] if purpose_col in param_data.columns else param_data
        
        # Plot polygon boundaries if available
        if polygon_gdf is not None and not polygon_gdf.empty:
            # Ensure same CRS
            if polygon_gdf.crs != param_data.crs:
                polygon_gdf = polygon_gdf.to_crs(param_data.crs)
            
            # Plot boundaries
            polygon_gdf.boundary.plot(ax=ax, color="black", linewidth=2)
            
            # Add plot type labels if available
            if "plot_type" in polygon_gdf.columns:
                for _, row in polygon_gdf.iterrows():
                    if row.geometry is not None and not row.geometry.is_empty:
                        centroid = row.geometry.centroid
                        label = str(row["plot_type"])[0].upper()
                        ax.text(
                            centroid.x, centroid.y, label,
                            ha="center", va="center",
                            fontsize=12, fontweight="bold", color="black",
                            bbox=dict(facecolor="white", alpha=0.7, boxstyle="circle,pad=0.2")
                        )
        
        # Plot points for this purpose
        if not purpose_data.empty:
            purpose_data.plot(
                ax=ax,
                column=value_col,
                cmap=cmap,
                markersize=point_size,
                alpha=0.8,
                vmin=vmin,
                vmax=vmax,
                legend=False,
                edgecolor="k",
                linewidth=0.5
            )
            ax.set_title(purpose, fontsize=14, fontweight="bold")
        else:
            ax.set_title(f"{purpose} (No Data)", fontsize=14, fontweight="bold")
        
        # Set equal aspect ratio and labels
        ax.set_aspect("equal", adjustable="box")
        ax.set_xlabel("Longitude", fontsize=10)
        ax.set_ylabel("Latitude", fontsize=10)
    
    # Synchronize axis limits
    if polygon_gdf is not None and not polygon_gdf.empty:
        # Use polygon bounds
        if polygon_gdf.crs != param_data.crs:
            polygon_gdf = polygon_gdf.to_crs(param_data.crs)
        xmin, ymin, xmax, ymax = polygon_gdf.total_bounds
        pad_x = (xmax - xmin) * 0.02
        pad_y = (ymax - ymin) * 0.02
        xmin, xmax = xmin - pad_x, xmax + pad_x
        ymin, ymax = ymin - pad_y, ymax + pad_y
    else:
        # Use data bounds
        xmin, ymin, xmax, ymax = param_data.total_bounds
        pad_x = (xmax - xmin) * 0.02
        pad_y = (ymax - ymin) * 0.02
        xmin, xmax = xmin - pad_x, xmax + pad_x
        ymin, ymax = ymin - pad_y, ymax + pad_y
    
    for ax in axes:
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(ymin, ymax)
    
    # Main title
    display_param_name = translate_parameter_for_display(param_name)
    fig.suptitle(f"Spatial Comparison - {display_param_name}", fontsize=16, fontweight="bold")
    
    # Add colorbar
    plt.subplots_adjust(right=0.85, top=0.9)
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
    sm._A = []
    cbar_ax = fig.add_axes([0.88, 0.15, 0.02, 0.7])
    cbar = fig.colorbar(sm, cax=cbar_ax)
    cbar.set_label(value_col, fontsize=12)
    
    plt.tight_layout()
    return fig

def create_box_plot(df, param_name, group_by_cols, param_col="Parâmetro", value_col="Resultado numérico"):
    """Create box plot for parameter grouped by specified columns"""
    if param_col not in df.columns or value_col not in df.columns:
        return None
    
    param_data = df[df[param_col] == param_name].copy()
    if param_data.empty:
        return None
    
    # Translate parameter name for display
    display_param_name = translate_parameter_for_display(param_name)
    
    # Handle multiple grouping columns
    if group_by_cols:
        # Filter to only include columns that exist in the data
        valid_group_cols = [col for col in group_by_cols if col in param_data.columns]
        
        if valid_group_cols:
            # Create a combined grouping column for multiple filters
            if len(valid_group_cols) == 1:
                combined_group_col = valid_group_cols[0]
            else:
                # Create a combined grouping column
                param_data['Combined_Group'] = param_data[valid_group_cols].apply(
                    lambda row: ' | '.join([f"{col}: {row[col]}" for col in valid_group_cols]), 
                    axis=1
                )
                combined_group_col = 'Combined_Group'
            
            fig = px.box(
                param_data,
                x=combined_group_col,
                y=value_col,
                color=combined_group_col,
                title=f"{t('distribution')} de {display_param_name} por {' + '.join(valid_group_cols)}",
                labels={value_col: "Valor"}
            )
        else:
            return None
    else:
        return None
    
    fig.update_layout(height=400)
    return fig

def create_statistical_summary(df, group_cols, param_col="Parâmetro", value_col="Resultado numérico"):
    """Create statistical summary grouped by specified columns"""
    if not all(col in df.columns for col in group_cols + [param_col, value_col]):
        return None
    
    # Calculate statistics
    stats = df.groupby(group_cols + [param_col])[value_col].agg([
        'count', 'mean', 'std', 'min', 'max', 'median'
    ]).round(3)
    
    lang = st.session_state.language
    if lang == 'en':
        stats.columns = ['N_Samples', 'Mean', 'Std_Deviation', 'Minimum', 'Maximum', 'Median']
    else:
        stats.columns = ['N_Amostras', 'Média', 'Desvio_Padrão', 'Mínimo', 'Máximo', 'Mediana']
    
    stats = stats.reset_index()
    return stats

def create_comparison_chart(df, param_name, group_col, stat_type="mean", param_col="Parâmetro", value_col="Resultado numérico"):
    """Create comparison chart showing statistics by group"""
    if param_col not in df.columns or group_col not in df.columns:
        return None
    
    param_data = df[df[param_col] == param_name].copy()
    if param_data.empty:
        return None
    
    # Calculate statistics by group
    if stat_type == "mean":
        stats = param_data.groupby(group_col)[value_col].mean()
        title_stat = STAT_NAMES[st.session_state.language]['mean']
    elif stat_type == "std":
        stats = param_data.groupby(group_col)[value_col].std()
        title_stat = STAT_NAMES[st.session_state.language]['std']
    else:
        stats = param_data.groupby(group_col)[value_col].median()
        title_stat = STAT_NAMES[st.session_state.language]['median']
    
    # Translate parameter name for display
    display_param_name = translate_parameter_for_display(param_name)
    
    fig = px.bar(
        x=stats.index,
        y=stats.values,
        title=f"{title_stat} de {display_param_name} por {group_col}",
        labels={"x": group_col, "y": f"{title_stat}"}
    )
    fig.update_layout(height=400)
    return fig

def create_overview_chart(df):
    """Create overview classification chart"""
    if "Classificação" not in df.columns:
        return None
    
    # Translate classifications for display
    df_display = df.copy()
    df_display["Classificação_Display"] = df_display["Classificação"].apply(translate_classification)
    
    summary = df_display["Classificação_Display"].value_counts()
    
    # Define colors for classifications
    color_map = {
        translate_classification("Muito Baixo"): "#FF4C4C",
        translate_classification("Baixo"): "#FFA04C", 
        translate_classification("Médio"): "#FFE14C",
        translate_classification("Bom"): "#9BEA8C",
        translate_classification("Muito Bom"): "#4CD964",
        translate_classification("Alto"): "#4CD964",
        translate_classification("Muito Alto"): "#1F7A1F"
    }
    
    colors = [color_map.get(cat, "#CCCCCC") for cat in summary.index]
    
    fig = go.Figure(data=[
        go.Bar(x=summary.index, y=summary.values, marker_color=colors)
    ])
    
    fig.update_layout(
        title=t('classification_summary'),
        xaxis_title=t('classification_breakdown').replace('📋 ', ''),
        yaxis_title=t('samples'),
        height=400
    )
    
    return fig

def generate_all_parameter_kriging_maps(df, points_gdf, polygon_gdf, classifier, param_col, value_col, 
                                      purpose_filter=None, depth_filter=None, grid_res=200):
    """Generate kriging maps for all parameters and return temporary image file paths"""
    if points_gdf is None or polygon_gdf is None:
        return []
    
    # Get unique parameters from the data
    unique_params = df[param_col].unique() if param_col in df.columns else []
    
    kriging_images = []
    temp_files = []
    
    for param in unique_params:
        try:
            # Check if parameter has valid classifications before creating map
            parameter_classifications = get_parameter_classifications(param)
            if not parameter_classifications:
                print(f"Skipping parameter '{param}' - no classification labels defined")
                continue
            
            # Create kriging map for this parameter
            kriging_result, bounds, xi, yi, error = create_kriging_map(
                points_gdf, polygon_gdf, param, classifier, param_col, value_col,
                purpose_filter, depth_filter, grid_res
            )
            
            if kriging_result is not None:
                # Calculate proper aspect ratio based on coordinate ranges
                xmin, ymin, xmax, ymax = bounds
                x_range = xmax - xmin
                y_range = ymax - ymin
                
                # Calculate aspect ratio to maintain proper proportions
                if y_range > 0:
                    aspect_ratio = x_range / y_range
                    # Limit aspect ratio to prevent extreme stretching
                    aspect_ratio = max(0.5, min(2.0, aspect_ratio))
                    fig_width = 8
                    fig_height = fig_width / aspect_ratio
                else:
                    fig_width, fig_height = 8, 6
                
                # Create figure with proper aspect ratio
                fig, ax = plt.subplots(figsize=(fig_width, fig_height))
                
                # Display the kriging map
                im = ax.imshow(kriging_result, origin="lower", extent=(xmin, xmax, ymin, ymax))
                
                # Add polygon boundary
                polygon_gdf.boundary.plot(ax=ax, edgecolor="black", linewidth=1.0)
                
                # Add plot type labels if available
                if 'plot_type' in polygon_gdf.columns:
                    for _, row in polygon_gdf.iterrows():
                        if not row.geometry.is_empty and not row.geometry.centroid.is_empty:
                            x_text, y_text = row.geometry.centroid.coords[0]
                            ax.text(x_text, y_text, row['plot_type'],
                                    ha='center', va='center', fontsize=14, 
                                    fontweight='bold', color='white',
                                    bbox=dict(boxstyle="round,pad=0.4", facecolor='black', alpha=0.8))
                
                # Set title and labels
                title = f"Kriging Map - {translate_parameter_for_display(param)}"
                if purpose_filter:
                    title += f" ({purpose_filter})"
                if depth_filter:
                    title += f" - Depth: {depth_filter}m"
                
                ax.set_title(title, fontsize=14, fontweight='bold', pad=12)
                ax.set_xlabel("Longitude", fontsize=12)
                ax.set_ylabel("Latitude", fontsize=12)
                
                # Add legend with classification levels, ranges, and units
                classification_colors = get_classification_colors()
                legend_handles = []
                legend_labels = []
                
                # Get parameter unit
                param_unit = get_parameter_unit(param)
                
                # Get only the classifications defined for this specific parameter
                parameter_classifications = get_parameter_classifications(param)
                
                for classification in parameter_classifications:
                    if classification in classification_colors:
                        color = classification_colors[classification]
                        translated_class = translate_classification(classification)
                        # Get threshold ranges for this parameter
                        threshold_ranges = get_parameter_thresholds(param, classification)
                        
                        if threshold_ranges:
                            legend_labels.append(f"{translated_class}: {threshold_ranges}")
                        else:
                            legend_labels.append(f"{translated_class}")
                        legend_handles.append(Patch(color=color, label=translated_class))
                
                if legend_handles:
                    # Create legend with detailed labels
                    legend = ax.legend(handles=legend_handles, labels=legend_labels, 
                                     loc='lower center', bbox_to_anchor=(0.5, -0.25), 
                                     ncol=min(3, len(legend_handles)), fontsize=14, 
                                     frameon=True, fancybox=True, shadow=True,
                                     title=f'Classification Levels ({param_unit})' if param_unit else 'Classification Levels',
                                     title_fontsize=16)
                    legend.get_frame().set_facecolor('white')
                    legend.get_frame().set_alpha(0.9)
                
                # Increase tick label font sizes
                ax.tick_params(axis='both', which='major', labelsize=10)
                
                # Adjust layout with more space for legend
                plt.tight_layout(pad=1.2)
                
                # Save to temporary file with higher DPI and proper bbox
                temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png', prefix=f'kriging_{param}_')
                plt.savefig(temp_img.name, dpi=300, bbox_inches='tight', facecolor='white', 
                           pad_inches=0.6)
                plt.close(fig)
                
                kriging_images.append({
                    'parameter': param,
                    'image_path': temp_img.name,
                    'title': title
                })
                temp_files.append(temp_img.name)
                
        except Exception as e:
            print(f"Error generating kriging map for {param}: {str(e)}")
            continue
    
    return kriging_images, temp_files

def generate_pdf_report(df, summary_stats, charts_data, project_name="Soil Analysis", 
                       points_gdf=None, polygon_gdf=None, classifier=None, 
                       param_col=None, value_col=None, purpose_filter=None, depth_filter=None):
    """Generate PDF report with kriging maps"""
    
    # Create temporary file
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    
    # Create PDF document
    doc = SimpleDocTemplate(
        temp_pdf.name,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        alignment=TA_LEFT,
        spaceAfter=12
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        alignment=TA_LEFT,
        spaceAfter=8
    )
    
    # Build story
    story = []
    
    # Title
    story.append(Paragraph(t('page_title'), title_style))
    story.append(Paragraph(f"Project: {project_name}", styles['Normal']))
    story.append(Paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Summary section
    story.append(Paragraph("Executive Summary", heading_style))
    
    summary_data = [
        ["Metric", "Value"],
        ["Total Samples", f"{summary_stats['total_samples']:,}"],
        ["Parameters Analyzed", f"{df['Parâmetro'].nunique() if 'Parâmetro' in df.columns else 'N/A'}"],
    ]
    
    # Add classification breakdown
    if 'classification_counts' in summary_stats:
        for classification, count in summary_stats['classification_counts'].items():
            percentage = summary_stats['classification_percentages'][classification]
            translated_class = translate_classification(classification)
            summary_data.append([f"{translated_class}", f"{count} ({percentage:.1f}%)"])
    
    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Add Kriging Maps section if geospatial data is available
    if points_gdf is not None and polygon_gdf is not None and classifier is not None:
        story.append(PageBreak())
        story.append(Paragraph("Kriging Maps", heading_style))
        story.append(Paragraph("Spatial interpolation maps showing the distribution of soil parameters across the field.", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Generate kriging maps for all parameters
        kriging_images, temp_files = generate_all_parameter_kriging_maps(
            df, points_gdf, polygon_gdf, classifier, param_col, value_col,
            purpose_filter, depth_filter
        )
        
        # Add each kriging map to the PDF
        for i, kriging_data in enumerate(kriging_images):
            try:
                # Add parameter title
                story.append(Paragraph(kriging_data['title'], subheading_style))
                
                # Add the image with dynamic sizing based on actual aspect ratio
                # Calculate aspect ratio from the saved image
                from PIL import Image as PILImage
                with PILImage.open(kriging_data['image_path']) as pil_img:
                    img_width, img_height = pil_img.size
                    aspect_ratio = img_width / img_height
                
                # Set maximum dimensions and maintain aspect ratio
                max_width = 7*inch
                max_height = 6*inch
                
                if aspect_ratio > 1:  # Landscape
                    width = min(max_width, max_height * aspect_ratio)
                    height = width / aspect_ratio
                else:  # Portrait
                    height = min(max_height, max_width / aspect_ratio)
                    width = height * aspect_ratio
                
                img = Image(kriging_data['image_path'], width=width, height=height)
                story.append(img)
                story.append(Spacer(1, 12))
                
                # Add page break after each map for better readability
                if i < len(kriging_images) - 1:
                    story.append(PageBreak())
                    
            except Exception as e:
                print(f"Error adding kriging map to PDF: {str(e)}")
                continue
    
    # Build PDF
    doc.build(story)
    
    # Clean up temporary image files
    if 'temp_files' in locals():
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except:
                pass
    
    # Read the PDF content
    with open(temp_pdf.name, 'rb') as f:
        pdf_content = f.read()
    
    # Clean up temporary PDF file
    try:
        os.unlink(temp_pdf.name)
    except:
        pass
    
    return pdf_content

def generate_docx_report(df, summary_stats, charts_data, project_name="Soil Analysis", 
                        points_gdf=None, polygon_gdf=None, classifier=None, 
                        param_col=None, value_col=None, purpose_filter=None, depth_filter=None):
    """Generate DOCX report with kriging maps"""
    
    # Create new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading(t('page_title'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project info
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()  # Empty line
    
    # Executive Summary section
    doc.add_heading('Executive Summary', level=1)
    
    # Summary table
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Add summary data
    summary_data = [
        ["Total Samples", f"{summary_stats['total_samples']:,}"],
        ["Parameters Analyzed", f"{df['Parâmetro'].nunique() if 'Parâmetro' in df.columns else 'N/A'}"],
    ]
    
    # Add classification breakdown
    if 'classification_counts' in summary_stats:
        for classification, count in summary_stats['classification_counts'].items():
            percentage = summary_stats['classification_percentages'][classification]
            translated_class = translate_classification(classification)
            summary_data.append([f"{translated_class}", f"{count} ({percentage:.1f}%)"])
    
    # Add data rows
    for metric, value in summary_data:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    doc.add_paragraph()  # Empty line
    
    # Add Kriging Maps section if geospatial data is available
    if points_gdf is not None and polygon_gdf is not None and classifier is not None:
        doc.add_heading('Kriging Maps', level=1)
        doc.add_paragraph('Spatial interpolation maps showing the distribution of soil parameters across the field.')
        doc.add_paragraph()  # Empty line
        
        # Generate kriging maps for all parameters
        kriging_images, temp_files = generate_all_parameter_kriging_maps(
            df, points_gdf, polygon_gdf, classifier, param_col, value_col,
            purpose_filter, depth_filter
        )
        
        # Add each kriging map to the DOCX
        for i, kriging_data in enumerate(kriging_images):
            try:
                # Add parameter title
                doc.add_heading(kriging_data['title'], level=2)
                
                # Add the image with appropriate sizing
                from PIL import Image as PILImage
                with PILImage.open(kriging_data['image_path']) as pil_img:
                    img_width, img_height = pil_img.size
                    aspect_ratio = img_width / img_height
                
                # Set maximum dimensions and maintain aspect ratio
                max_width = 6.5  # inches
                max_height = 5.0  # inches
                
                if aspect_ratio > 1:  # Landscape
                    width = min(max_width, max_height * aspect_ratio)
                    height = width / aspect_ratio
                else:  # Portrait
                    height = min(max_height, max_width / aspect_ratio)
                    width = height * aspect_ratio
                
                # Add image to document
                doc.add_picture(kriging_data['image_path'], width=Inches(width))
                
                # Add page break after each map for better readability
                if i < len(kriging_images) - 1:
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"Error adding kriging map to DOCX: {str(e)}")
                continue
    
    # Save to BytesIO buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Clean up temporary image files
    if 'temp_files' in locals():
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except:
                pass
    
    return doc_buffer.getvalue()

def detect_geometry_columns(df):
    """Detect potential geometry columns in the dataframe"""
    geometry_columns = []
    
    for col in df.columns:
        col_lower = col.lower()
        # Check for common geometry column names
        if any(keyword in col_lower for keyword in ['geometry', 'geom', 'shape', 'wkb', 'wkt', 'coordinates', 'coord']):
            geometry_columns.append(col)
        else:
            # Check if column contains WKB or WKT data
            sample_values = df[col].dropna().head(10)
            if len(sample_values) > 0:
                # Check if values look like WKB (hex strings or binary)
                sample_str = str(sample_values.iloc[0])
                if (len(sample_str) > 20 and 
                    (all(c in '0123456789ABCDEFabcdef' for c in sample_str.replace(' ', '')) or
                     isinstance(sample_values.iloc[0], bytes))):
                    geometry_columns.append(col)
                # Check if values look like WKT
                elif any(wkt_keyword in sample_str.upper() for wkt_keyword in ['POINT', 'POLYGON', 'LINESTRING', 'MULTIPOINT', 'MULTIPOLYGON']):
                    geometry_columns.append(col)
    
    return geometry_columns

def convert_wkb_to_geometry(wkb_data):
    """Convert WKB data to Shapely geometry objects"""
    try:
        if isinstance(wkb_data, str):
            # Try as hex string first
            try:
                return wkb.loads(wkb_data, hex=True)
            except:
                # Try as regular WKB
                try:
                    return wkb.loads(bytes.fromhex(wkb_data))
                except:
                    # Try as WKT
                    return wkt.loads(wkb_data)
        elif isinstance(wkb_data, bytes):
            return wkb.loads(wkb_data)
        else:
            # Try to convert to string and parse as WKT
            return wkt.loads(str(wkb_data))
    except Exception as e:
        print(f"Error converting WKB data: {e}")
        return None

def create_geodataframe_from_geometry(df, geometry_col, crs='EPSG:4326'):
    """Create a GeoDataFrame from a dataframe with geometry column"""
    try:
        # Create a copy of the dataframe
        gdf_data = df.copy()
        
        # Convert geometry column to Shapely geometries
        geometries = []
        for idx, geom_data in enumerate(gdf_data[geometry_col]):
            if pd.isna(geom_data):
                geometries.append(None)
            else:
                geom = convert_wkb_to_geometry(geom_data)
                if geom is not None:
                    geometries.append(geom)
                else:
                    geometries.append(None)
        
        # Replace the geometry column with converted geometries
        gdf_data[geometry_col] = geometries
        
        # Remove rows with invalid geometries
        gdf_data = gdf_data.dropna(subset=[geometry_col])
        
        # Create GeoDataFrame
        gdf = gpd.GeoDataFrame(gdf_data, geometry=geometry_col, crs=crs)
        
        return gdf
        
    except Exception as e:
        print(f"Error creating GeoDataFrame: {e}")
        return None

def get_parameter_translations():
    """Get parameter translations from English to Portuguese"""
    return {
        "aluminum saturation": "Saturação por alumínio (m%)",
        "aluminum soil": "Acidez trocável (Al3+)",
        "base saturation": "Saturação por bases (V%)",
        "calcium soil": "Cálcio trocável (Ca2+)",
        "copper soil": "Cobre (Cu)",
        "effective cation exchange capacity": "CTC efetiva (t)",
        "iron soil": "Ferro (Fe)",
        "magnesium soil": "Magnésio trocável (Mg2+)",
        "manganese soil": "Manganês (Mn)",
        "organic matter": "Matéria orgânica (MO)",
        "ph in cacl2": "pH em CaCl₂",
        "ph in water": "pH em H₂O",
        "phosphorus soil": "P - disponível",
        "potassium soil": "Potássio trocável (K+)",
        "potential acidity": "Acidez potencial (H+Al)",
        "sulfur soil": "Enxofre (S)",
        "total cation exchange capacity at ph 7.00": "CTC a pH 7,0 (T)",
        "total organic carbon soil": "Carbono orgânico total",
        "zinc soil": "Zinco (Zn)",
    }

def get_parameter_unit(param):
    """Get the unit for a parameter"""
    unit_mapping = {
        "Cálcio trocável (Ca2+)": "cmolc/dm³",
        "Magnésio trocável (Mg2+)": "cmolc/dm³", 
        "Potássio trocável (K+)": "mg/dm³",
        "Matéria orgânica (MO)": "dag/kg",
        "P - disponível": "mg/dm³",
        "Acidez trocável (Al3+)": "cmolc/dm³",
        "Acidez potencial (H+Al)": "cmolc/dm³",
        "CTC efetiva (t)": "cmolc/dm³",
        "CTC a pH 7,0 (T)": "cmolc/dm³",
        "pH em CaCl₂": "",
        "pH em H₂O": "",
        "Saturação por bases (V%)": "%",
        "Saturação por alumínio (m%)": "%",
        "Cobre (Cu)": "mg/dm³",
        "Ferro (Fe)": "mg/dm³",
        "Manganês (Mn)": "mg/dm³",
        "Zinco (Zn)": "mg/dm³",
        "Enxofre (S)": "mg/dm³",
        "Carbono orgânico total": "g/kg",
        "Boro (B)": "mg/dm³",
        "P disponível (Resina, Sequeiro)": "mg/dm³",
        "P disponível (Resina, Irrigado)": "mg/dm³",
        # Add English variations
        "calcium soil": "cmolc/dm³",
        "magnesium soil": "cmolc/dm³",
        "potassium soil": "mg/dm³",
        "organic matter": "dag/kg",
        "phosphorus soil": "mg/dm³",
        "copper soil": "mg/dm³",
        "iron soil": "mg/dm³",
        "manganese soil": "mg/dm³",
        "zinc soil": "mg/dm³",
        "sulfur soil": "mg/dm³",
        "boron soil": "mg/dm³"
    }
    
    # Try exact match first
    if param in unit_mapping:
        return unit_mapping[param]
    
    # Try case-insensitive match
    param_lower = param.lower()
    for key, unit in unit_mapping.items():
        if param_lower in key.lower() or key.lower() in param_lower:
            return unit
    
    return ""

def get_parameter_classifications(param):
    """Get the specific classifications defined for a parameter"""
    # Get all parameters from both MB and BA defaults
    all_params = {**PARAMS_MB_DEFAULT, **PARAMS_BA_DEFAULT}
    
    # Create a mapping for common parameter variations
    param_mapping = {
        "calcium soil": "Cálcio trocável (Ca2+)",
        "magnesium soil": "Magnésio trocável (Mg2+)",
        "potassium soil": "Potássio trocável (K+)",
        "organic matter": "Matéria orgânica (MO)",
        "phosphorus soil": "P - disponível",
        "copper soil": "Cobre (Cu)",
        "iron soil": "Ferro (Fe)",
        "manganese soil": "Manganês (Mn)",
        "zinc soil": "Zinco (Zn)",
        "sulfur soil": "Enxofre (S)",
        "aluminum soil": "Acidez trocável (Al3+)",
        "potential acidity": "Acidez potencial (H+Al)",
        "effective cation exchange capacity": "CTC efetiva (t)",
        "total cation exchange capacity at ph 7.00": "CTC a pH 7,0 (T)",
        "base saturation": "Saturação por bases (V%)",
        "aluminum saturation": "Saturação por alumínio (m%)",
        "ph in cacl2": "pH em CaCl₂",
        "ph in water": "pH em H₂O",
        "boron soil": "Boro (B)"
    }
    
    # Try to find the correct parameter name
    target_param = param
    if param.lower() in param_mapping:
        target_param = param_mapping[param.lower()]
    
    # Try exact match first
    if target_param in all_params:
        return [classification for _, _, classification in all_params[target_param]]
    else:
        # Try fuzzy matching
        param_lower = param.lower()
        for available_param in all_params.keys():
            if param_lower in available_param.lower() or available_param.lower() in param_lower:
                return [classification for _, _, classification in all_params[available_param]]
    
    return []

def get_parameter_thresholds(param, classification):
    """Get threshold ranges for a parameter and classification"""
    # Get all parameters from both MB and BA defaults
    all_params = {**PARAMS_MB_DEFAULT, **PARAMS_BA_DEFAULT}
    
    # Create a mapping for common parameter variations
    param_mapping = {
        "calcium soil": "Cálcio trocável (Ca2+)",
        "magnesium soil": "Magnésio trocável (Mg2+)",
        "potassium soil": "Potássio trocável (K+)",
        "organic matter": "Matéria orgânica (MO)",
        "phosphorus soil": "P - disponível",
        "copper soil": "Cobre (Cu)",
        "iron soil": "Ferro (Fe)",
        "manganese soil": "Manganês (Mn)",
        "zinc soil": "Zinco (Zn)",
        "sulfur soil": "Enxofre (S)",
        "aluminum soil": "Acidez trocável (Al3+)",
        "potential acidity": "Acidez potencial (H+Al)",
        "effective cation exchange capacity": "CTC efetiva (t)",
        "total cation exchange capacity at ph 7.00": "CTC a pH 7,0 (T)",
        "base saturation": "Saturação por bases (V%)",
        "aluminum saturation": "Saturação por alumínio (m%)",
        "ph in cacl2": "pH em CaCl₂",
        "ph in water": "pH em H₂O",
        "boron soil": "Boro (B)"
    }
    
    # Try to find the correct parameter name
    target_param = param
    if param.lower() in param_mapping:
        target_param = param_mapping[param.lower()]
    
    # Try exact match first
    if target_param in all_params:
        thresholds = all_params[target_param]
        for min_val, max_val, class_name in thresholds:
            if class_name == classification:
                if max_val == float('inf'):
                    return f"> {min_val:.1f}"
                elif min_val == 0 and max_val < 1:
                    return f"≤ {max_val:.2f}"
                elif min_val == 0:
                    return f"≤ {max_val:.1f}"
                else:
                    return f"{min_val:.2f} - {max_val:.1f}"
    else:
        # Try fuzzy matching
        param_lower = param.lower()
        for available_param in all_params.keys():
            if (param_lower in available_param.lower() or 
                available_param.lower() in param_lower or
                any(keyword in param_lower for keyword in ["calcium", "magnesium", "potassium", "copper", "iron", "zinc", "sulfur", "organic", "phosphorus", "boron"])):
                thresholds = all_params[available_param]
                for min_val, max_val, class_name in thresholds:
                    if class_name == classification:
                        if max_val == float('inf'):
                            return f"> {min_val:.1f}"
                        elif min_val == 0 and max_val < 1:
                            return f"≤ {max_val:.2f}"
                        elif min_val == 0:
                            return f"≤ {max_val:.1f}"
                        else:
                            return f"{min_val:.2f} - {max_val:.1f}"
    return ""

def get_column_translations():
    """Get column name translations from English to Portuguese"""
    return {
        "plot_type": "Tratamento",
        "sampling_plan_purpose": "Data amostragem",
        "depth_range_top_m": "profundidade superior",
        "depth_range_bottom_m": "profundidade inferior",
        "campo_sample_number": "ID-LAB-CAMPO",
        "translated_standard_parameter": "Parâmetro",
        "numeric_result": "Resultado numérico",
        "unit_pad": "Unidade",
        "geometry": "geometria",
    }

def detect_and_translate_english_data(df):
    """
    Detect if the dataframe contains English column names and translate only column names to Portuguese.
    Parameter values are kept in their original language for classification.
    """
    column_translations = get_column_translations()
    
    # Check column names for English patterns
    english_columns_found = any(col in column_translations for col in df.columns)
    
    if english_columns_found:
        st.info("🔄 English column names detected! Translating column names to Portuguese...")
        
        # Translate only column names, keep parameter values in original language
        df_translated = df.rename(columns=column_translations)
        
        # Add missing columns with default values if needed
        expected_columns = [
            "Tratamento", "Data amostragem", "profundidade superior", 
            "profundidade inferior", "ID-LAB-CAMPO", "Parâmetro", 
            "Resultado numérico", "Unidade", "geometria"
        ]
        
        for col in expected_columns:
            if col not in df_translated.columns:
                if col == "geometria":
                    df_translated[col] = "POINT(0 0)"  # Default geometry
                elif col == "Unidade":
                    df_translated[col] = "unit"  # Default unit
                elif col == "Data amostragem":
                    df_translated[col] = "2024-01-01"  # Default date
                else:
                    df_translated[col] = "unknown"  # Default value
        
        st.success("✅ Column translation completed! Parameter values kept in original language for classification.")
        return df_translated
    
    return df

def translate_classification_to_english(portuguese_classification):
    """Translate Portuguese classifications to English"""
    translation_map = {
        "Muito Baixo": "Very Low",
        "Baixo": "Low", 
        "Médio": "Medium",
        "Bom": "Good",
        "Muito Bom": "Very Good",
        "Alto": "High",
        "Muito Alto": "Very High",
        "Classificação não definida": "Classification not defined",
        "Não classificado": "Not classified"
    }
    return translation_map.get(portuguese_classification, portuguese_classification)

def translate_parameter_to_english(portuguese_parameter):
    """Translate Portuguese parameter names back to original English names"""
    # Reverse mapping of the parameter translations
    parameter_translation_map = {
        "Saturação por alumínio (m%)": "aluminum saturation",
        "Acidez trocável (Al3+)": "aluminum soil",
        "Saturação por bases (V%)": "base saturation",
        "Cálcio trocável (Ca2+)": "calcium soil",
        "Cobre (Cu)": "copper soil",
        "CTC efetiva (t)": "effective cation exchange capacity",
        "Ferro (Fe)": "iron soil",
        "Magnésio trocável (Mg2+)": "magnesium soil",
        "Manganês (Mn)": "manganese soil",
        "Matéria orgânica (MO)": "organic matter",
        "pH em CaCl₂": "ph in cacl2",
        "pH em H₂O": "ph in water",
        "P - disponível": "phosphorus soil",
        "Potássio trocável (K+)": "potassium soil",
        "Acidez potencial (H+Al)": "potential acidity",
        "Enxofre (S)": "sulfur soil",
        "CTC a pH 7,0 (T)": "total cation exchange capacity at ph 7.00",
        "Carbono orgânico total": "total organic carbon soil",
        "Zinco (Zn)": "zinc soil",
    }
    return parameter_translation_map.get(portuguese_parameter, portuguese_parameter)

def translate_column_names_to_english(df):
    """Translate Portuguese column names back to original English names"""
    column_translation_map = {
        "Tratamento": "plot_type",
        "Data amostragem": "sampling_plan_purpose", 
        "profundidade superior": "depth_range_top_m",
        "profundidade inferior": "depth_range_bottom_m",
        "ID-LAB-CAMPO": "campo_sample_number",
        "Parâmetro": "translated_standard_parameter",
        "Resultado numérico": "numeric_result",
        "Unidade": "unit_pad",
        "geometria": "geometry",
        "Classificação": "Classification",
        "Classificação_Média": "Mean_Classification",
        "N_Amostras": "N_Samples",
        "Média": "Mean",
        "Mediana": "Median", 
        "Desvio_Padrão": "Std_Deviation",
        "Mínimo": "Minimum",
        "Máximo": "Maximum"
    }
    
    # Rename columns
    df_translated = df.rename(columns=column_translation_map)
    
    # Translate classification column values if it exists
    classification_cols = ["Classification", "Mean_Classification"]
    for col in classification_cols:
        if col in df_translated.columns:
            df_translated[col] = df_translated[col].apply(translate_classification_to_english)
    
    # Translate parameter names back to English if the parameter column exists
    parameter_cols = ["translated_standard_parameter"]
    for col in parameter_cols:
        if col in df_translated.columns:
            df_translated[col] = df_translated[col].apply(translate_parameter_to_english)
    
    return df_translated

def apply_excel_colors(worksheet, classification_col_name="Classificação"):
    """Apply color formatting to the classification column in Excel"""
    
    # Color palette matching the original specification
    color_fills = {
        "Muito Baixo": PatternFill(start_color="FF4C4C", end_color="FF4C4C", fill_type="solid"),  # vermelho
        "Baixo":       PatternFill(start_color="FFA04C", end_color="FFA04C", fill_type="solid"),  # laranja
        "Médio":       PatternFill(start_color="FFE14C", end_color="FFE14C", fill_type="solid"),  # amarelo
        "Bom":         PatternFill(start_color="9BEA8C", end_color="9BEA8C", fill_type="solid"),  # verde claro
        "Muito Bom":   PatternFill(start_color="4CD964", end_color="4CD964", fill_type="solid"),  # verde médio
        "Alto":        PatternFill(start_color="4CD964", end_color="4CD964", fill_type="solid"),  # verde médio
        "Muito Alto":  PatternFill(start_color="1F7A1F", end_color="1F7A1F", fill_type="solid"),  # verde escuro
        
        # English translations
        "Very Low": PatternFill(start_color="FF4C4C", end_color="FF4C4C", fill_type="solid"),
        "Low":      PatternFill(start_color="FFA04C", end_color="FFA04C", fill_type="solid"),
        "Medium":   PatternFill(start_color="FFE14C", end_color="FFE14C", fill_type="solid"),
        "Good":     PatternFill(start_color="9BEA8C", end_color="9BEA8C", fill_type="solid"),
        "Very Good": PatternFill(start_color="4CD964", end_color="4CD964", fill_type="solid"),
        "High":     PatternFill(start_color="4CD964", end_color="4CD964", fill_type="solid"),
        "Very High": PatternFill(start_color="1F7A1F", end_color="1F7A1F", fill_type="solid"),
        
        # Additional variations and non-classified values
        "Classificação não definida": PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid"),  # cinza claro
        "Não classificado": PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid"),  # cinza claro
    }
    
    # Find the classification column
    header = {}
    for idx, cell in enumerate(worksheet[1], start=1):
        if cell.value:
            header[cell.value] = idx
    
    classification_col = header.get(classification_col_name)
    if not classification_col:
        return  # Column not found
    
    # Apply colors to classification cells
    for row_num in range(2, worksheet.max_row + 1):  # Start from row 2 (skip header)
        cell = worksheet.cell(row=row_num, column=classification_col)
        classification_value = str(cell.value).strip() if cell.value else ""
        
        if classification_value in color_fills:
            cell.fill = color_fills[classification_value]

def create_medias_sheet(writer, classified_df, param_col, value_col, language='pt'):
    """
    Create medias (averages) sheet by plot_type with color-coded classifications
    
    Args:
        writer: Excel writer object
        classified_df: DataFrame with classified data
        param_col: Parameter column name (user-selected, could be 'translated_standard_parameter' or other)
        value_col: Value column name (user-selected, could be 'numeric_result' or other)
        language: Language for output ('pt' or 'en')
    """
    try:
        # Get the actual column names based on the data language
        # Check what columns exist in the dataframe
        plot_type_col = None
        depth_col = None
        
        # Try to find plot_type column (could be 'Tratamento' in PT or 'plot_type' in EN)
        if 'Tratamento' in classified_df.columns:
            plot_type_col = 'Tratamento'
        elif 'plot_type' in classified_df.columns:
            plot_type_col = 'plot_type'
            
        # Try to find depth column (could be 'profundidade inferior' in PT or 'depth_range_bottom_m' in EN)
        if 'profundidade inferior' in classified_df.columns:
            depth_col = 'profundidade inferior'
        elif 'depth_range_bottom_m' in classified_df.columns:
            depth_col = 'depth_range_bottom_m'
        
        # Check if required columns exist
        required_cols = [plot_type_col, 'sampling_plan_purpose', depth_col]
        missing_cols = [col for col in required_cols if col is None or col not in classified_df.columns]
        if missing_cols or param_col not in classified_df.columns or value_col not in classified_df.columns:
            print(f"Missing columns for medias sheet: {missing_cols}")
            print(f"Available columns: {list(classified_df.columns)}")
            return
        
        # Group by parameter, plot_type, sampling_plan_purpose, and depth, calculate mean and classification
        medias_data = []
        
        for param in classified_df[param_col].unique():
            param_data = classified_df[classified_df[param_col] == param]
            
            for plot_type in param_data[plot_type_col].unique():
                plot_data = param_data[param_data[plot_type_col] == plot_type]
                
                for purpose in plot_data['sampling_plan_purpose'].unique():
                    purpose_data = plot_data[plot_data['sampling_plan_purpose'] == purpose]
                    
                    for depth in purpose_data[depth_col].unique():
                        depth_data = purpose_data[purpose_data[depth_col] == depth]
                        
                        if not depth_data.empty:
                            mean_value = depth_data[value_col].mean()
                            
                            # Get classification for the mean value using the original parameter name
                            classifier = SoilClassifier()
                            # Use the parameter name as it appears in the data (could be English or Portuguese)
                            classification = classifier.classify_value(param, mean_value)
                            
                            # Translate classification for display
                            translated_classification = translate_classification(classification)
                            
                            if language == 'en':
                                medias_data.append({
                                    'Parameter': param,
                                    'Plot_Type': plot_type,
                                    'Sampling_Purpose': purpose,
                                    'Depth': depth,
                                    'Mean_Value': round(mean_value, 3),
                                    'Classification': translated_classification
                                })
                            else:
                                medias_data.append({
                                    'Parâmetro': param,
                                    'Tipo_Plot': plot_type,
                                    'Propósito_Amostragem': purpose,
                                    'Profundidade': depth,
                                    'Valor_Médio': round(mean_value, 3),
                                    'Classificação': translated_classification
                                })
        
        if medias_data:
            medias_df = pd.DataFrame(medias_data)
            
            # Sort by parameter, plot_type, purpose, and depth
            if language == 'en':
                medias_df = medias_df.sort_values(['Parameter', 'Plot_Type', 'Sampling_Purpose', 'Depth'])
                classification_col = 'Classification'
                sheet_name = 'Means'
            else:
                medias_df = medias_df.sort_values(['Parâmetro', 'Tipo_Plot', 'Propósito_Amostragem', 'Profundidade'])
                classification_col = 'Classificação'
                sheet_name = 'Médias'
            
            # Write to Excel
            medias_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Apply colors to the classification column
            apply_excel_colors(writer.book[sheet_name], classification_col)
            
            print(f"Created medias sheet with {len(medias_data)} rows")
        else:
            print("No data found for medias sheet")
            
    except Exception as e:
        print(f"Error creating medias sheet: {e}")

def create_color_legend_sheet(workbook, language='pt'):
    """Create a color legend sheet in the workbook"""
    
    # Define sheet names and headers based on language
    if language == 'en':
        sheet_name = "Color_Legend"
        header1 = "Classification"
        header2 = "Color"
        color_definitions = [
            ("Very Low", "FF4C4C"),
            ("Low", "FFA04C"),
            ("Medium", "FFE14C"),
            ("Good", "9BEA8C"),
            ("Very Good", "4CD964"),
            ("High", "4CD964"),
            ("Very High", "1F7A1F")
        ]
    else:
        sheet_name = "Legenda_Cores"
        header1 = "Classificação"
        header2 = "Cor"
        color_definitions = [
            ("Muito Baixo", "FF4C4C"),
            ("Baixo", "FFA04C"),
            ("Médio", "FFE14C"),
            ("Bom", "9BEA8C"),
            ("Muito Bom", "4CD964"),
            ("Alto", "4CD964"),
            ("Muito Alto", "1F7A1F")
        ]
    
    # Remove existing legend sheet if it exists
    for sheet_name_to_check in ["Legenda_Cores", "Color_Legend"]:
        if sheet_name_to_check in workbook.sheetnames:
            workbook.remove(workbook[sheet_name_to_check])
    
    # Create new legend sheet
    legend_sheet = workbook.create_sheet(sheet_name)
    
    # Headers
    legend_sheet.cell(row=1, column=1, value=header1)
    legend_sheet.cell(row=1, column=2, value=header2)
    
    # Add legend entries
    for row_idx, (classification, color_code) in enumerate(color_definitions, start=2):
        # Classification name
        legend_sheet.cell(row=row_idx, column=1, value=classification)
        
        # Color cell
        color_cell = legend_sheet.cell(row=row_idx, column=2, value="")
        color_cell.fill = PatternFill(start_color=color_code, end_color=color_code, fill_type="solid")
    
    # Adjust column widths
    legend_sheet.column_dimensions['A'].width = 15
    legend_sheet.column_dimensions['B'].width = 10

def kriging_interpolation(x, y, z, xmin, xmax, ymin, ymax, grid_res=500, variogram_model='linear'):
    """Perform kriging interpolation over custom extent - exact same as notebook"""
    xi = np.linspace(xmin, xmax, grid_res)
    yi = np.linspace(ymin, ymax, grid_res)
    xi, yi = np.meshgrid(xi, yi)

    OK = OrdinaryKriging(
        x, y, z,
        variogram_model=variogram_model,
        verbose=False,
        enable_plotting=False
    )

    zi, _ = OK.execute('grid', xi[0], yi[:, 0])
    return xi, yi, zi

def create_kriging_map(points_gdf, polygon_gdf, parameter_name, classifier, param_col, value_col, 
                      purpose_filter=None, depth_filter=None, grid_res=500):
    """Create kriging map for a specific parameter"""
    try:
        # Check if parameter has valid classifications defined
        parameter_classifications = get_parameter_classifications(parameter_name)
        if not parameter_classifications:
            return None, f"No classification labels defined for parameter: {parameter_name}"
        
        # Filter points data
        filtered_points = points_gdf.copy()
        
        if param_col in filtered_points.columns:
            filtered_points = filtered_points[filtered_points[param_col] == parameter_name]
        
        if purpose_filter and 'sampling_plan_purpose' in filtered_points.columns:
            filtered_points = filtered_points[filtered_points['sampling_plan_purpose'] == purpose_filter]
            
        if depth_filter and 'depth_range_bottom_m' in filtered_points.columns:
            filtered_points = filtered_points[filtered_points['depth_range_bottom_m'] == depth_filter]
        
        if filtered_points.empty:
            return None, "No data points found for the selected criteria"
        
        # Get coordinates and values
        x = filtered_points.geometry.x.values
        y = filtered_points.geometry.y.values
        z = filtered_points[value_col].values
        
        # Get polygon bounds
        xmin, ymin, xmax, ymax = polygon_gdf.total_bounds
        
        # Perform kriging with higher resolution for better smoothing
        xi, yi, zi = kriging_interpolation(x, y, z, xmin, xmax, ymin, ymax, grid_res=grid_res)
        
        # Classify interpolated values
        classify_vectorized = np.vectorize(lambda val: classifier.classify_value(parameter_name, val))
        categories = classify_vectorized(zi)
        
        # Get classification colors from the app
        classification_colors = get_classification_colors()
        
        # Create RGB image using the same technique as the notebook
        rgb_image = np.zeros((*zi.shape, 3), dtype=np.uint8)
        for label in np.unique(categories):
            if label in classification_colors:
                mask = categories == label
                hex_color = classification_colors[label]
                color = tuple(int(hex_color[i:i+2], 16) for i in (1, 3, 5))
                rgb_image[mask] = color
            else:
                # Default white for unclassified (like in notebook)
                mask = categories == label
                rgb_image[mask] = (255, 255, 255)
        
        # Apply polygon mask to set areas outside polygon to white (like in notebook)
        from shapely.geometry import mapping
        from rasterio.features import geometry_mask
        from affine import Affine
        
        shapes = [mapping(geom) for geom in polygon_gdf.geometry]
        pixel_width = xi[0, 1] - xi[0, 0]
        pixel_height = yi[1, 0] - yi[0, 0]
        affine = Affine(pixel_width, 0, xmin, 0, pixel_height, ymin)
        
        mask = geometry_mask(geometries=shapes, transform=affine, invert=True, out_shape=zi.shape)
        rgb_image[~mask] = 255  # White outside polygon
        
        return rgb_image, (xmin, ymin, xmax, ymax), xi, yi, None
        
    except Exception as e:
        return None, f"Error creating kriging map: {str(e)}"

def get_classification_colors():
    """Get classification colors for visualization"""
    return {
        "Muito Baixo": "#070707",
        "Baixo": "#d7191c", 
        "Médio": "#ffa849",
        "Bom": "#abdda4",
        "Muito Bom": "#2b83ba",
        "Alto": "#4CD964",
        "Muito Alto": "#1F7A1F",
        "Classificação não definida": "#E0E0E0",
        "Valor inválido": "#E0E0E0"
    }

def create_comprehensive_statistics_with_classification(df, grouping_cols, param_col, value_col, language='pt'):
    """
    Create comprehensive statistics grouped by specified columns with mean classification
    """
    # Check if all grouping columns exist
    existing_cols = [col for col in grouping_cols if col in df.columns]
    if not existing_cols:
        return None
    
    # Add parameter column to grouping
    group_cols = existing_cols + [param_col]
    
    try:
        # Calculate comprehensive statistics
        stats_df = df.groupby(group_cols)[value_col].agg([
            'count',   # N_Amostras / N_Samples
            'mean',    # Média / Mean
            'median',  # Mediana / Median
            'std',     # Desvio_Padrão / Std_Deviation
            'min',     # Mínimo / Minimum
            'max',     # Máximo / Maximum
        ]).round(4).reset_index()
        
        # Rename columns based on language
        if language == 'en':
            stats_df.columns = existing_cols + [param_col, 'N_Samples', 'Mean', 'Median', 'Std_Deviation', 'Minimum', 'Maximum']
            classification_col = 'Mean_Classification'
        else:
            stats_df.columns = existing_cols + [param_col, 'N_Amostras', 'Média', 'Mediana', 'Desvio_Padrão', 'Mínimo', 'Máximo']
            classification_col = 'Classificação_Média'
        
        # Apply classification to the mean values
        classifier = SoilClassifier()
        mean_col = 'Mean' if language == 'en' else 'Média'
        stats_df[classification_col] = stats_df.apply(
            lambda row: classifier.classify_value(row[param_col], row[mean_col]), 
            axis=1
        )
        
        # Translate classifications if English
        if language == 'en':
            stats_df[classification_col] = stats_df[classification_col].apply(translate_classification_to_english)
            
            # Also translate the column names to English (including grouping columns)
            stats_df = translate_column_names_to_english(stats_df)
        
        # Sort by grouping columns and parameter
        sort_cols = list(stats_df.columns[:len(existing_cols)]) + [list(stats_df.columns)[len(existing_cols)]]  # First few columns + parameter column
        stats_df = stats_df.sort_values(sort_cols)
        
        return stats_df
        
    except Exception as e:
        print(f"Error creating comprehensive statistics: {e}")
        return None

def main():
    # Language selector in sidebar
    st.sidebar.selectbox(
        t('language'),
        options=['pt', 'en'],
        format_func=lambda x: '🇧🇷 Português' if x == 'pt' else '🇺🇸 English',
        key='language'
    )
    
    st.title(t('main_title'))
    st.markdown(t('subtitle'))
    
    # Sidebar
    st.sidebar.header(t('settings'))
    
    # Database connection section
    st.sidebar.subheader("🗄️ Database Connection")
    use_database = st.sidebar.checkbox("Load data from database", help="Check this to load data directly from the database instead of uploading files")
    
    if use_database:
        # Database configuration
        st.sidebar.markdown("**Database Configuration:**")
        
        st.sidebar.info("ℹ️ Using agbenefits pipeline database connection method")
        st.sidebar.markdown("**Requirements:**")
        st.sidebar.markdown("- Active gcloud session (same as agbenefits pipeline)")
        st.sidebar.markdown("- utils/db.py file in project directory")
        st.sidebar.markdown("- Environment variables: DB_HOST, DB_NAME, DB_USER")
        
        # Check if .env file exists
        env_file_exists = os.path.exists('.env')
        if not env_file_exists:
            st.sidebar.warning("⚠️ .env file not found")
            st.sidebar.markdown("**Create a .env file with:**")
            st.sidebar.code("""
DB_HOST=your_host
DB_NAME=your_database
DB_USER=your_username
DB_PORT=5432
            """)
        else:
            st.sidebar.success("✅ .env file found")
        
        # Try to setup database connection using agbenefits pipeline method
        db_connection_ok = setup_database_connection()
        
        if db_connection_ok:
            # Field ID input
            field_id = st.sidebar.number_input(
                "Field ID",
                min_value=1,
                value=1,
                help="Enter the field ID to retrieve data for"
            )
            
            # Load data button
            if st.sidebar.button("Load Data from Database", type="primary"):
                with st.spinner("Loading data from database..."):
                    # Retrieve soil samples
                    soil_samples_df = retrieve_soil_samples_from_db(field_id)
                    
                    if not soil_samples_df.empty:
                        st.success(f"✅ Loaded {len(soil_samples_df)} soil samples from database")
                        
                        # Store raw data in session state for column selection
                        st.session_state['raw_db_data'] = soil_samples_df
                        st.session_state['data_source'] = 'database'
                        st.session_state['field_id'] = field_id
                        st.session_state['db_data_loaded'] = True
                    else:
                        st.error("❌ No soil samples found for the specified field ID")
            
            # Show column selection if database data is loaded
            if st.session_state.get('db_data_loaded', False) and 'raw_db_data' in st.session_state:
                st.sidebar.markdown("**📊 Column Selection for Database Data:**")
                
                # Get all columns from the loaded data
                soil_samples_df = st.session_state['raw_db_data']
                all_columns = list(soil_samples_df.columns)
                numeric_columns = list(soil_samples_df.select_dtypes(include=['number']).columns)
                
                # Parameter column selection
                param_col = st.sidebar.selectbox(
                    "Select Parameter Column",
                    options=all_columns,
                    index=0 if all_columns else None,
                    help="Select the column that contains parameter names (e.g., 'parameter', 'analysis_type')"
                )
                
                # Result column selection
                result_col = st.sidebar.selectbox(
                    "Select Result Column",
                    options=numeric_columns,
                    index=0 if numeric_columns else None,
                    help="Select the column that contains numeric results (e.g., 'result', 'value', 'concentration')"
                )
                
                # Geometry column selection
                geom_col = st.sidebar.selectbox(
                    "Select Geometry Column",
                    options=all_columns,
                    index=0 if all_columns else None,
                    help="Select the column that contains geometry data (WKB, WKT, or coordinates)"
                )
                
                # CRS input for geometry
                crs_input = st.sidebar.text_input(
                    "CRS (Coordinate Reference System)",
                    value="EPSG:4326",
                    help="Enter the CRS for the geometry data (e.g., EPSG:4326, EPSG:3857)"
                )
                
                # Show data preview
                with st.sidebar.expander("📋 Data Preview", expanded=False):
                    st.write("**Available columns:**")
                    for col in all_columns:
                        st.write(f"- {col}")
                    
                    st.write(f"\n**Data shape:** {soil_samples_df.shape[0]} rows × {soil_samples_df.shape[1]} columns")
                    
                    st.write("**First few rows:**")
                    st.dataframe(soil_samples_df.head(3), use_container_width=True)
                
                # Process button
                if st.sidebar.button("Process Database Data", type="primary"):
                    if param_col and result_col and geom_col:
                        # Process the data similar to file upload
                        try:
                            # Create a processed dataframe similar to file upload
                            processed_df = soil_samples_df.copy()
                            
                            # Rename columns to match expected format
                            processed_df = processed_df.rename(columns={
                                param_col: 'Parameter',
                                result_col: 'Result',
                                geom_col: 'geometry'
                            })
                            
                            # Create geometry column
                            points_gdf = create_geodataframe_from_geometry(processed_df, 'geometry', crs_input)
                            
                            if points_gdf is not None:
                                # Convert geometry to string for session state compatibility
                                gdf_for_session = points_gdf.copy()
                                if 'geometry' in gdf_for_session.columns:
                                    gdf_for_session['geometry'] = gdf_for_session['geometry'].astype(str)
                                
                                # Store processed data in session state
                                st.session_state['uploaded_data'] = gdf_for_session
                                st.session_state['data_source'] = 'database'
                                st.session_state['field_id'] = field_id
                                st.session_state['points_gdf'] = points_gdf
                                
                                st.success("✅ Database data processed successfully!")
                                st.info(f"Processed {len(points_gdf)} records with geometry")
                                
                                # Try to retrieve field boundaries
                                boundary_gdf = retrieve_field_boundaries_from_db(field_id)
                                if boundary_gdf is not None:
                                    st.success("✅ Loaded field boundaries from database")
                                    # Save boundary to temporary file for compatibility
                                    temp_boundary_file = tempfile.NamedTemporaryFile(delete=False, suffix='.geojson')
                                    temp_boundary_file.close()  # Close the file handle first
                                    boundary_gdf.to_file(temp_boundary_file.name, driver='GeoJSON')
                                    # Store the temporary file path and GeoDataFrame
                                    st.session_state['polygon_file'] = temp_boundary_file
                                    st.session_state['polygon_gdf'] = boundary_gdf
                                    st.info(f"Loaded {len(boundary_gdf)} polygon boundaries")
                                else:
                                    st.info("ℹ️ No field boundaries found in database")
                            else:
                                st.error("❌ Failed to process geometry data")
                                
                        except Exception as e:
                            st.error(f"❌ Error processing data: {str(e)}")
                    else:
                        st.error("❌ Please select all required columns")
    
    # File upload (only show if not using database)
    if not use_database:
        st.sidebar.subheader("📁 File Upload")
    else:
        st.sidebar.subheader("📁 Alternative: File Upload")
    
    # File upload
    uploaded_file = st.sidebar.file_uploader(
        t('upload_file'),
        type=['xlsx', 'csv'],
        help=t('upload_help')
    )
    
    # Geometry column selection for kriging maps (will be shown after file upload)
    geometry_cols = None
    points_geometry_col = None
    polygon_geometry_col = None
    
    # Polygon geometry upload
    polygon_file = st.sidebar.file_uploader(
        "Upload Polygon Geometry",
        type=['geojson', 'shp', 'gpkg'],
        help="Upload GeoJSON, Shapefile, or GeoPackage containing field boundaries/polygons"
    )
    
    # Store polygon file in session state
    if polygon_file is not None:
        st.session_state['polygon_file'] = polygon_file
    else:
        st.session_state['polygon_file'] = None
    
    # Project name input
    project_name = st.sidebar.text_input(
        t('project_name'),
        value=t('default_project'),
        help=t('project_help')
    )
    
    # Check if data is available (either from file upload or database)
    data_available = uploaded_file is not None or st.session_state.get('uploaded_data') is not None
    
    if data_available:
        try:
            # Determine data source and load accordingly
            if uploaded_file is not None:
                # Read file
                if uploaded_file.name.endswith('.xlsx'):
                    df = pd.read_excel(uploaded_file)
                else:
                    df = pd.read_csv(uploaded_file)
            else:
                # Use database data
                df = st.session_state['uploaded_data']
            
            # Keep original column names for user mapping
            # df = detect_and_translate_english_data(df)  # Removed automatic translation
            
            st.success(f"{t('file_loaded')} {len(df)} {t('lines_found')}")
            
            # Show data preview
            with st.expander(t('data_preview'), expanded=False):
                st.dataframe(df.head(10))
                st.info(f"{t('columns_available')} {', '.join(df.columns.tolist())}")
            
            # Column mapping (only for file uploads, database data is already processed)
            if uploaded_file is not None:
                st.sidebar.subheader(t('column_mapping'))
                
                param_col = st.sidebar.selectbox(
                    t('parameter_column'),
                    options=df.columns.tolist(),
                    index=0 if len(df.columns) > 0 else None,
                    help=t('parameter_help')
                )
                
                value_col = st.sidebar.selectbox(
                    t('value_column'),
                    options=df.columns.tolist(),
                    index=1 if len(df.columns) > 1 else 0,
                    help=t('value_help')
                )
            else:
                # For database data, use the already processed column names
                param_col = 'Parameter'
                value_col = 'Result'
            
            # Initialize classifier
            classifier = SoilClassifier()
            
            # Custom parameters section
            with st.sidebar.expander(t('custom_parameters'), expanded=False):
                st.markdown(t('add_custom'))
                
                custom_param = st.text_input(t('parameter_name'))
                param_type = st.selectbox(t('parameter_type'), ["MB", "BA"])
                
                if st.button(t('add_parameter')):
                    if custom_param:
                        # Default ranges for new parameter
                        default_ranges = [(0, 20, "Muito Baixo"), (20, 40, "Baixo"), (40, 60, "Médio"), (60, 80, "Bom"), (80, float('inf'), "Muito Bom")]
                        classifier.add_custom_parameter(custom_param, default_ranges, param_type)
                        st.success(t('parameter_added').format(custom_param))
            
            # Process classification
            if st.sidebar.button(t('run_classification'), type="primary"):
                with st.spinner(t('processing')):
                    
                    # Classify data
                    classified_df = classifier.classify_dataframe(df, param_col, value_col)
                    
                    # Calculate summary
                    summary_stats = create_classification_summary(classified_df)
                    
                    # Store in session state
                    st.session_state['classified_df'] = classified_df
                    st.session_state['summary_stats'] = summary_stats
                    st.session_state['project_name'] = project_name
                    
                    # Detect geometry columns (only for file uploads)
                    if uploaded_file is not None:
                        geometry_cols = detect_geometry_columns(classified_df)
                        st.session_state['geometry_cols'] = geometry_cols
                    else:
                        # For database data, geometry is already processed
                        st.session_state['geometry_cols'] = []
                
                st.success(t('classification_completed'))
            
            # Display results if available
            if 'classified_df' in st.session_state:
                classified_df = st.session_state['classified_df']
                summary_stats = st.session_state['summary_stats']
                geometry_cols = st.session_state.get('geometry_cols', [])
                
                # Geometry column selection in sidebar
                if geometry_cols or st.session_state.get('points_gdf') is not None:
                    st.sidebar.subheader("🗺️ Points Geometry Selection")
                    
                    if uploaded_file is not None and geometry_cols:
                        # Points geometry column (from Excel file)
                        points_geometry_col = st.sidebar.selectbox(
                            "Select Points Geometry Column",
                            options=['None'] + geometry_cols,
                            help="Select the column containing point geometries (soil sample locations) from your Excel file"
                        )
                    else:
                        # For database data, geometry is already processed
                        points_geometry_col = 'geometry'
                        st.info("✅ Points geometry loaded from database")
                    
                    # CRS selection
                    crs_input = st.sidebar.text_input(
                        "Coordinate Reference System (CRS)",
                        value="EPSG:4326",
                        help="Enter the CRS code (e.g., EPSG:4326 for WGS84)"
                    )
                    
                    # Store selections in session state
                    st.session_state['points_geometry_col'] = points_geometry_col if points_geometry_col != 'None' else None
                    st.session_state['crs_input'] = crs_input
                else:
                    st.sidebar.info("ℹ️ No geometry columns detected in the uploaded file. Kriging maps will not be available.")
                
                # Main content area
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric(
                        t('total_samples'),
                        summary_stats['total_samples']
                    )
                
                with col2:
                    if 'classification_counts' in summary_stats:
                        most_common = max(summary_stats['classification_counts'], key=summary_stats['classification_counts'].get)
                        st.metric(
                            t('predominant_class'),
                            translate_classification(most_common),
                            f"{summary_stats['classification_percentages'][most_common]:.1f}%"
                        )
                
                with col3:
                    parameters_count = classified_df[param_col].nunique()
                    st.metric(
                        t('unique_parameters'),
                        parameters_count
                    )
                
                # Charts
                st.subheader(t('visualizations'))
                
                # Overview chart
                overview_fig = create_overview_chart(classified_df)
                if overview_fig:
                    st.plotly_chart(overview_fig, use_container_width=True)
                
                # Advanced Analysis Section
                st.subheader(t('advanced_analysis'))
                
                # Parameter selection
                if param_col in classified_df.columns:
                    unique_params = classified_df[param_col].unique()
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        selected_param = st.selectbox(
                            t('select_parameter'),
                            options=unique_params
                        )
                    
                    with col2:
                        # Group by column selection with display names (multiselect)
                        display_to_actual = get_grouping_columns_with_display_names(
                            classified_df, 
                            [param_col, value_col, "Classificação"]
                        )
                        
                        display_options = list(display_to_actual.keys())
                        selected_display_names = st.multiselect(
                            t('group_by') + " (Multiple Selection)",
                            options=display_options,
                            help=t('group_help') + " - You can select multiple columns to aggregate by multiple filters",
                            default=[]
                        )
                        
                        # Get actual column names
                        group_by_cols = [display_to_actual.get(display_name, display_name) for display_name in selected_display_names]
                    
                    if selected_param:
                        # Create tabs for different analysis types
                        tab1, tab2, tab3, tab4, tab5 = st.tabs([
                            t('distribution'), 
                            t('box_plot'), 
                            t('statistics'), 
                            t('comparison'),
                            "🗺️ Spatial Plots"
                        ])
                        
                        with tab1:
                            # Distribution chart options
                            col1, col2 = st.columns([3, 1])
                            with col2:
                                separate_by_classification = st.checkbox(
                                    "Separate by Classification",
                                    value=False,
                                    help="Show each classification level in separate subplots"
                                )
                            
                            # Distribution chart
                            if group_by_cols:
                                param_fig = create_parameter_chart(classified_df, selected_param, group_by_cols, param_col, value_col, separate_by_classification)
                            else:
                                param_fig = create_parameter_chart(classified_df, selected_param, None, param_col, value_col, separate_by_classification)
                            
                            if param_fig:
                                st.plotly_chart(param_fig, use_container_width=True)
                        
                        with tab2:
                            # Box plot
                            if group_by_cols:
                                box_fig = create_box_plot(classified_df, selected_param, group_by_cols, param_col, value_col)
                                if box_fig:
                                    st.plotly_chart(box_fig, use_container_width=True)
                                else:
                                    st.info(t('box_plot_unavailable'))
                            else:
                                st.info("Please select at least one grouping column to create box plots")
                        
                        with tab3:
                            # Statistical summary
                            st.markdown(f"#### {t('statistical_summary')}")
                            
                            # Use the same grouping columns as selected for plots
                            if group_by_cols:
                                stats_df = create_statistical_summary(
                                    classified_df, 
                                    group_by_cols, 
                                    param_col, 
                                    value_col
                                )
                                
                                if stats_df is not None:
                                    # Filter for selected parameter
                                    param_stats = stats_df[stats_df[param_col] == selected_param]
                                    if not param_stats.empty:
                                        st.dataframe(param_stats, use_container_width=True)
                                        
                                        # Download stats
                                        csv = param_stats.to_csv(index=False)
                                        st.download_button(
                                            label=t('download_stats'),
                                            data=csv,
                                            file_name=f"estatisticas_{selected_param}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                            mime="text/csv"
                                        )
                                    else:
                                        st.info(t('no_stats_available'))
                                else:
                                    st.error("Failed to generate statistical summary")
                            else:
                                st.info("Please select grouping columns to view statistical summary")
                        
                        with tab4:
                            # Comparison charts
                            st.markdown(f"#### {t('comparison_between_groups')}")
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                # Get grouping columns for comparison
                                display_to_actual_comp = get_grouping_columns_with_display_names(
                                    classified_df, 
                                    [param_col, value_col, "Classificação"]
                                )
                                
                                selected_comparison_display = st.selectbox(
                                    t('compare_by'),
                                    options=list(display_to_actual_comp.keys()),
                                    key="comparison_col"
                                )
                                
                                # Get actual column name
                                comparison_col = display_to_actual_comp.get(selected_comparison_display, selected_comparison_display)
                            
                            with col2:
                                stat_type = st.selectbox(
                                    t('statistic'),
                                    options=["mean", "std", "median"],
                                    format_func=lambda x: STAT_NAMES[st.session_state.language][x]
                                )
                            
                            if comparison_col:
                                comparison_fig = create_comparison_chart(
                                    classified_df, 
                                    selected_param, 
                                    comparison_col, 
                                    stat_type,
                                    param_col,
                                    value_col
                                )
                                if comparison_fig:
                                    st.plotly_chart(comparison_fig, use_container_width=True)
                                
                                # Show numerical comparison
                                param_data = classified_df[classified_df[param_col] == selected_param]
                                if not param_data.empty and comparison_col in param_data.columns:
                                    comparison_stats = param_data.groupby(comparison_col)[value_col].agg([
                                        'count', 'mean', 'std', 'median'
                                    ]).round(3)
                                    
                                    if st.session_state.language == 'en':
                                        comparison_stats.columns = ['N', 'Mean', 'Std Dev', 'Median']
                                    else:
                                        comparison_stats.columns = ['N', 'Média', 'Desvio Padrão', 'Mediana']
                                    
                                    st.markdown(f"##### {t('numerical_values')}")
                                    st.dataframe(comparison_stats, use_container_width=True)
                
                # Classification breakdown
                st.subheader(t('classification_breakdown'))
                
                if 'classification_counts' in summary_stats:
                    breakdown_cols = st.columns(len(summary_stats['classification_counts']))
                    
                    for i, (classification, count) in enumerate(summary_stats['classification_counts'].items()):
                        with breakdown_cols[i]:
                            percentage = summary_stats['classification_percentages'][classification]
                            st.metric(
                                translate_classification(classification),
                                f"{count}",
                                f"{percentage:.1f}%"
                            )
                
                # Complete Statistical Overview
                st.subheader(t('statistical_overview'))
                
                with st.expander(t('complete_statistics'), expanded=False):
                    # Allow user to select multiple grouping columns with display names
                    display_to_actual_complete = get_grouping_columns_with_display_names(
                        classified_df, 
                        [param_col, value_col, "Classificação"]
                    )
                    
                    if display_to_actual_complete:
                        # Determine default selection
                        default_options = []
                        if "plot_type" in classified_df.columns:
                            default_options = [translate_column_for_display("plot_type")]
                        elif list(display_to_actual_complete.keys()):
                            default_options = [list(display_to_actual_complete.keys())[0]]
                        
                        selected_stat_display_groups = st.multiselect(
                            t('select_stat_groups'),
                            options=list(display_to_actual_complete.keys()),
                            default=default_options,
                            help=t('stat_groups_help')
                        )
                        
                        # Convert display names back to actual column names
                        selected_stat_groups = [display_to_actual_complete[display_name] for display_name in selected_stat_display_groups]
                        
                        if selected_stat_groups:
                            # Calculate comprehensive statistics
                            complete_stats = create_statistical_summary(
                                classified_df, 
                                selected_stat_groups, 
                                param_col, 
                                value_col
                            )
                            
                            if complete_stats is not None:
                                st.markdown(f"#### {t('complete_stats_table')}")
                                st.dataframe(complete_stats, use_container_width=True)
                                
                                # Download complete stats
                                csv_complete = complete_stats.to_csv(index=False)
                                st.download_button(
                                    label=t('download_complete_stats'),
                                    data=csv_complete,
                                    file_name=f"estatisticas_completas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                    mime="text/csv"
                                )
                                
                                # Show summary by parameter
                                st.markdown(f"#### {t('parameter_summary')}")
                                
                                col_map = {'N_Amostras': 'sum', 'Média': 'mean', 'Desvio_Padrão': 'mean', 'Mínimo': 'min', 'Máximo': 'max'} if st.session_state.language == 'pt' else {'N_Samples': 'sum', 'Mean': 'mean', 'Std_Deviation': 'mean', 'Minimum': 'min', 'Maximum': 'max'}
                                
                                param_summary = complete_stats.groupby(param_col).agg(col_map).round(3)
                                
                                st.dataframe(param_summary, use_container_width=True)
                    else:
                        st.info(t('no_grouping_columns'))
                
                # Data table
                with st.expander(t('classified_data'), expanded=False):
                    # Translate classifications for display
                    display_df = classified_df.copy()
                    display_df["Classificação"] = display_df["Classificação"].apply(translate_classification)
                    st.dataframe(display_df, use_container_width=True)
                        
                    with tab5:
                        # Spatial Plots Section
                        st.markdown("#### 🗺️ Spatial Visualization")
                        
                        # Check if we have spatial data
                        points_geometry_col = st.session_state.get('points_geometry_col')
                        polygon_file = st.session_state.get('polygon_file', None)
                        
                        # For database data, check if we have processed geometry
                        has_spatial_data = False
                        if uploaded_file is not None:
                            # File upload: check geometry column selection and polygon file
                            has_spatial_data = points_geometry_col and polygon_file is not None
                        else:
                            # Database data: check if we have processed geometry data
                            has_spatial_data = (st.session_state.get('points_gdf') is not None and 
                                              (st.session_state.get('polygon_file') is not None or 
                                               st.session_state.get('polygon_gdf') is not None))
                        
                        if has_spatial_data:
                            try:
                                # Create GeoDataFrames
                                if uploaded_file is not None:
                                    # File upload: create from geometry column
                                    points_gdf = create_geodataframe_from_geometry(
                                        classified_df, points_geometry_col, 
                                        st.session_state.get('crs_input', 'EPSG:4326')
                                    )
                                else:
                                    # Database data: use already processed geometry
                                    points_gdf = st.session_state.get('points_gdf')
                                
                                # Load polygon from uploaded file or database
                                if uploaded_file is not None:
                                    # File upload: load from uploaded file
                                    if polygon_file is not None:
                                        if polygon_file.name.endswith('.geojson'):
                                            polygon_gdf = gpd.read_file(polygon_file)
                                        elif polygon_file.name.endswith('.shp'):
                                            polygon_gdf = gpd.read_file(polygon_file)
                                        elif polygon_file.name.endswith('.gpkg'):
                                            polygon_gdf = gpd.read_file(polygon_file)
                                        else:
                                            polygon_gdf = None
                                    else:
                                        polygon_gdf = None
                                else:
                                    # Database data: use stored GeoDataFrame or load from temporary file
                                    polygon_gdf = st.session_state.get('polygon_gdf')
                                    if polygon_gdf is None:
                                        # Fallback to temporary file if GeoDataFrame not available
                                        current_polygon_file = st.session_state.get('polygon_file')
                                        if current_polygon_file is not None and hasattr(current_polygon_file, 'name'):
                                            polygon_gdf = gpd.read_file(current_polygon_file.name)
                                
                                if points_gdf is not None and polygon_gdf is not None:
                                    # Spatial plot controls
                                    col1, col2 = st.columns([2, 1])
                                    
                                    with col1:
                                        # Plot type selection
                                        plot_type = st.selectbox(
                                            "Plot Type",
                                            options=["Single Plot", "Comparison Plot"],
                                            help="Choose between single plot or comparison plot with two panels"
                                        )
                                    
                                    with col2:
                                        # Color scheme
                                        cmap = st.selectbox(
                                            "Color Scheme",
                                            options=["viridis", "plasma", "inferno", "magma", "cividis", "RdYlBu_r"],
                                            index=0,
                                            help="Choose color scheme for the plot"
                                        )
                                    
                                    # Filter controls
                                    col3, col4 = st.columns(2)
                                    
                                    with col3:
                                        # Purpose filter
                                        purpose_options = ["All"]
                                        if 'sampling_plan_purpose' in points_gdf.columns:
                                            purpose_options.extend(points_gdf['sampling_plan_purpose'].unique())
                                        selected_purpose = st.selectbox(
                                            "Sampling Purpose Filter:",
                                            options=purpose_options,
                                            key="spatial_purpose"
                                        )
                                    
                                    with col4:
                                        # Depth filter
                                        depth_options = ["All"]
                                        if 'depth_range_bottom_m' in points_gdf.columns:
                                            depth_options.extend(sorted(points_gdf['depth_range_bottom_m'].unique()))
                                        selected_depth = st.selectbox(
                                            "Depth Filter:",
                                            options=depth_options,
                                            key="spatial_depth"
                                        )
                                    
                                    # Generate spatial plot
                                    if st.button("Generate Spatial Plot", type="primary"):
                                        with st.spinner("Generating spatial plot..."):
                                            try:
                                                if plot_type == "Single Plot":
                                                    spatial_fig = create_spatial_plot(
                                                        points_gdf, selected_param,
                                                        polygon_gdf=polygon_gdf,
                                                        purpose_filter=selected_purpose if selected_purpose != "All" else None,
                                                        depth_filter=selected_depth if selected_depth != "All" else None,
                                                        param_col=param_col,
                                                        value_col=value_col,
                                                        cmap=cmap
                                                    )
                                                else:  # Comparison Plot
                                                    # Get available purposes for comparison
                                                    available_purposes = []
                                                    if 'sampling_plan_purpose' in points_gdf.columns:
                                                        available_purposes = list(points_gdf['sampling_plan_purpose'].unique())
                                                    
                                                    if len(available_purposes) >= 2:
                                                        spatial_fig = create_spatial_comparison_plot(
                                                            points_gdf, selected_param,
                                                            polygon_gdf=polygon_gdf,
                                                            purposes=(available_purposes[0], available_purposes[1]),
                                                            param_col=param_col,
                                                            value_col=value_col,
                                                            cmap=cmap
                                                        )
                                                    else:
                                                        st.warning("Need at least 2 different sampling purposes for comparison plot")
                                                        spatial_fig = None
                                                
                                                if spatial_fig:
                                                    st.pyplot(spatial_fig, use_container_width=True)
                                                    plt.close(spatial_fig)  # Close to free memory
                                                else:
                                                    st.error("Failed to generate spatial plot. Check your data and filters.")
                                                    
                                            except Exception as e:
                                                st.error(f"Error generating spatial plot: {str(e)}")
                                
                            except Exception as e:
                                st.error(f"Error loading spatial data: {str(e)}")
                        else:
                            if uploaded_file is not None:
                                # File upload mode
                                if not points_geometry_col:
                                    st.info("ℹ️ Please select a points geometry column in the sidebar to enable spatial plots.")
                                if not polygon_file:
                                    st.info("ℹ️ Please upload a polygon geometry file in the sidebar to enable spatial plots.")
                            else:
                                # Database mode
                                if st.session_state.get('points_gdf') is None:
                                    st.info("ℹ️ Please process the database data with geometry columns to enable spatial plots.")
                                if (st.session_state.get('polygon_file') is None and 
                                    st.session_state.get('polygon_gdf') is None):
                                    st.info("ℹ️ Please process the database data to load polygon boundaries for spatial plots.")
            
            # Kriging Maps Section
            points_geometry_col = st.session_state.get('points_geometry_col')
            crs_input = st.session_state.get('crs_input', 'EPSG:4326')
            
            # Check if we have spatial data for kriging maps
            has_kriging_data = False
            if uploaded_file is not None:
                # File upload: check geometry column selection and polygon file
                has_kriging_data = points_geometry_col and polygon_file is not None
            else:
                # Database data: check if we have processed geometry data
                has_kriging_data = (st.session_state.get('points_gdf') is not None and 
                                  (st.session_state.get('polygon_file') is not None or 
                                   st.session_state.get('polygon_gdf') is not None))
            
            if has_kriging_data:
                st.subheader("🗺️ Kriging Maps")
                
                try:
                    # Create points GeoDataFrame
                    if uploaded_file is not None:
                        # File upload: create from geometry column
                        points_gdf = create_geodataframe_from_geometry(classified_df, points_geometry_col, crs_input)
                    else:
                        # Database data: use already processed geometry
                        points_gdf = st.session_state.get('points_gdf')
                    
                    # Load polygon GeoDataFrame from uploaded file or database
                    if uploaded_file is not None:
                        # File upload: load from uploaded file
                        if polygon_file is not None:
                            if polygon_file.name.endswith('.geojson'):
                                polygon_gdf = gpd.read_file(polygon_file)
                            elif polygon_file.name.endswith('.shp'):
                                polygon_gdf = gpd.read_file(polygon_file)
                            elif polygon_file.name.endswith('.gpkg'):
                                polygon_gdf = gpd.read_file(polygon_file)
                            else:
                                polygon_gdf = None
                        else:
                            polygon_gdf = None
                    else:
                        # Database data: use stored GeoDataFrame or load from temporary file
                        polygon_gdf = st.session_state.get('polygon_gdf')
                        if polygon_gdf is None:
                            # Fallback to temporary file if GeoDataFrame not available
                            current_polygon_file = st.session_state.get('polygon_file')
                            if current_polygon_file is not None and hasattr(current_polygon_file, 'name'):
                                polygon_gdf = gpd.read_file(current_polygon_file.name)
                    
                    if points_gdf is None or polygon_gdf is None:
                        st.error("❌ Failed to load geometry data")
                    else:
                        # Ensure both GeoDataFrames have the same CRS
                        polygon_gdf = polygon_gdf.to_crs(points_gdf.crs)
                        
                        st.success("✅ Geometry data loaded successfully!")
                    
                    # Map configuration
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        # Parameter selection - only show parameters with valid classifications
                        if param_col in points_gdf.columns:
                            all_params = points_gdf[param_col].unique()
                            # Filter to only include parameters that have classification labels defined
                            available_params = [param for param in all_params if get_parameter_classifications(param)]
                            
                            if available_params:
                                selected_param_map = st.selectbox(
                                    "Select Parameter for Map:",
                                    options=available_params,
                                    key="map_param"
                                )
                            else:
                                st.warning("No parameters with valid classification labels found in the data")
                                selected_param_map = None
                                
                                # Show which parameters were found but don't have classifications
                                params_without_classifications = [param for param in all_params if not get_parameter_classifications(param)]
                                if params_without_classifications:
                                    st.info(f"Parameters found but without classification labels: {', '.join(params_without_classifications)}")
                        else:
                            st.error(f"Parameter column '{param_col}' not found in points data")
                            selected_param_map = None
                    
                    with col2:
                        # Purpose filter
                        purpose_options = ["All"]
                        if 'sampling_plan_purpose' in points_gdf.columns:
                            purpose_options.extend(points_gdf['sampling_plan_purpose'].unique())
                        selected_purpose = st.selectbox(
                            "Sampling Purpose:",
                            options=purpose_options,
                            key="map_purpose"
                        )
                        purpose_filter = None if selected_purpose == "All" else selected_purpose
                    
                    with col3:
                        # Depth filter
                        depth_options = ["All"]
                        if 'depth_range_bottom_m' in points_gdf.columns:
                            depth_options.extend(sorted(points_gdf['depth_range_bottom_m'].unique()))
                        selected_depth = st.selectbox(
                            "Depth:",
                            options=depth_options,
                            key="map_depth"
                        )
                        depth_filter = None if selected_depth == "All" else selected_depth
                    
                    # Grid resolution with better defaults for display
                    grid_res = st.slider("Grid Resolution:", min_value=100, max_value=500, value=500, step=25,
                                        help="Higher resolution = more detailed but slower. Recommended: 300-500 for good balance")
                    
                    if selected_param_map and st.button("🗺️ Generate Kriging Map", type="primary"):
                        with st.spinner("Generating kriging map..."):
                            # Create kriging map
                            kriging_result, bounds, xi, yi, error = create_kriging_map(
                                points_gdf, polygon_gdf, selected_param_map, classifier,
                                param_col, value_col, purpose_filter, depth_filter, grid_res
                            )
                            
                            if kriging_result is not None:
                                # Create matplotlib figure with compact sizing similar to reference
                                fig, ax = plt.subplots(figsize=(8, 5))
                                
                                # Display the kriging map
                                xmin, ymin, xmax, ymax = bounds
                                ax.imshow(kriging_result, origin="lower", extent=(xmin, xmax, ymin, ymax))
                                
                                # Add polygon boundary
                                polygon_gdf.boundary.plot(ax=ax, edgecolor="black", linewidth=1.2)
                                
                                # Add plot type labels if available
                                if 'plot_type' in polygon_gdf.columns:
                                    for _, row in polygon_gdf.iterrows():
                                        if not row.geometry.is_empty and not row.geometry.centroid.is_empty:
                                            x_text, y_text = row.geometry.centroid.coords[0]
                                            ax.text(x_text, y_text, row['plot_type'],
                                                    ha='center', va='center', fontsize=6, 
                                                    fontweight='bold', color='white',
                                                    bbox=dict(boxstyle="round,pad=0.1", facecolor='black', alpha=0.7))
                                
                                # Set title and labels - more compact
                                title = f"Kriging Map - {translate_parameter_for_display(selected_param_map)}"
                                if purpose_filter:
                                    title += f" ({purpose_filter})"
                                if depth_filter:
                                    title += f" - Depth: {depth_filter}m"
                                
                                ax.set_title(title, fontsize=10, fontweight='bold', pad=8)
                                ax.set_xlabel("Longitude", fontsize=8)
                                ax.set_ylabel("Latitude", fontsize=8)
                                
                                # Adjust layout to prevent overflow
                                plt.tight_layout(pad=0.5)
                                
                                # Display the map
                                st.pyplot(fig, use_container_width=True, clear_figure=True)
                                
                                # Create legend separately below the map
                                classification_colors = get_classification_colors()
                                legend_handles = []
                                legend_labels = []
                                
                                # Get only the classifications defined for this specific parameter
                                parameter_classifications = get_parameter_classifications(selected_param_map)
                                
                                for classification in parameter_classifications:
                                    if classification in classification_colors:
                                        color = classification_colors[classification]
                                        translated_class = translate_classification(classification)
                                        threshold_ranges = get_parameter_thresholds(selected_param_map, classification)
                                        
                                        if threshold_ranges:
                                            legend_labels.append(f"{translated_class}: {threshold_ranges}")
                                        else:
                                            legend_labels.append(f"{translated_class}")
                                        legend_handles.append(Patch(color=color, label=translated_class))
                                
                                if legend_handles:
                                    # Create legend as a separate element below the map
                                    legend_fig, legend_ax = plt.subplots(figsize=(12, 1))
                                    legend_ax.legend(handles=legend_handles, loc='center', ncol=len(legend_handles),
                                                   fontsize=9, frameon=False)
                                    legend_ax.axis('off')
                                    plt.tight_layout()
                                    st.pyplot(legend_fig, use_container_width=True, clear_figure=True)
                                
                                # Download option - recreate figure for download with compact layout
                                fig_download, ax_download = plt.subplots(figsize=(8, 5))
                                ax_download.imshow(kriging_result, origin="lower", extent=(xmin, xmax, ymin, ymax))
                                polygon_gdf.boundary.plot(ax=ax_download, edgecolor="black", linewidth=1.2)
                                
                                if 'plot_type' in polygon_gdf.columns:
                                    for _, row in polygon_gdf.iterrows():
                                        if not row.geometry.is_empty and not row.geometry.centroid.is_empty:
                                            x_text, y_text = row.geometry.centroid.coords[0]
                                            ax_download.text(x_text, y_text, row['plot_type'],
                                                            ha='center', va='center', fontsize=6, 
                                                            fontweight='bold', color='white',
                                                            bbox=dict(boxstyle="round,pad=0.1", facecolor='black', alpha=0.7))
                                
                                ax_download.set_title(title, fontsize=10, fontweight='bold', pad=8)
                                ax_download.set_xlabel("Longitude", fontsize=8)
                                ax_download.set_ylabel("Latitude", fontsize=8)
                                
                                if legend_handles:
                                    # Create legend below the map for download
                                    ax_download.legend(handles=legend_handles, loc='lower center', bbox_to_anchor=(0.5, -0.15), 
                                                     ncol=len(legend_handles), fontsize=8, frameon=False)
                                
                                plt.tight_layout(pad=0.5)
                                
                                img_buffer = io.BytesIO()
                                fig_download.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
                                img_buffer.seek(0)
                                plt.close(fig_download)  # Close to free memory
                                
                                st.download_button(
                                    label="📥 Download Kriging Map (PNG)",
                                    data=img_buffer.getvalue(),
                                    file_name=f"kriging_map_{selected_param_map}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                                    mime="image/png"
                                )
                                
                            else:
                                st.error(f"❌ {error}")
                    
                    # Show data preview
                    with st.expander("📊 GeoJSON Data Preview", expanded=False):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.subheader("Points Data")
                            st.write(f"**Shape:** {points_gdf.shape}")
                            st.write(f"**Columns:** {', '.join(points_gdf.columns.tolist())}")
                            if len(points_gdf) > 0:
                                st.dataframe(points_gdf.head())
                        
                        with col2:
                            st.subheader("Polygon Data")
                            st.write(f"**Shape:** {polygon_gdf.shape}")
                            st.write(f"**Columns:** {', '.join(polygon_gdf.columns.tolist())}")
                            if len(polygon_gdf) > 0:
                                st.dataframe(polygon_gdf.head())
                
                except Exception as e:
                    st.error(f"❌ Error processing geometry data: {str(e)}")
                    st.info("Please ensure your points geometry column contains valid WKB/WKT data and your polygon file is a valid GeoJSON, Shapefile, or GeoPackage.")
            else:
                if uploaded_file is not None:
                    # File upload mode
                    if not points_geometry_col:
                        st.info("ℹ️ Please select a points geometry column from your Excel file to enable kriging maps.")
                    if not polygon_file:
                        st.info("ℹ️ Please upload a polygon geometry file (GeoJSON, Shapefile, or GeoPackage) to enable kriging maps.")
                else:
                    # Database mode
                    if st.session_state.get('points_gdf') is None:
                        st.info("ℹ️ Please process the database data with geometry columns to enable kriging maps.")
                    if (st.session_state.get('polygon_file') is None and 
                        st.session_state.get('polygon_gdf') is None):
                        st.info("ℹ️ Please process the database data to load polygon boundaries for kriging maps.")
            
            # Report Generation
            st.subheader(t('pdf_report_generation'))
            
            col1, col2, col3 = st.columns([1, 1, 1])
            
            with col1:
                if st.button(t('generate_pdf'), type="primary"):
                    with st.spinner(t('generating_pdf')):
                        try:
                            # Pass geospatial data and other parameters for kriging maps
                            pdf_content = generate_pdf_report(
                                classified_df, 
                                summary_stats, 
                                None,  # charts_data placeholder
                                st.session_state.get('project_name', t('default_project')),
                                points_gdf=points_gdf if 'points_gdf' in locals() else None,
                                polygon_gdf=polygon_gdf if 'polygon_gdf' in locals() else None,
                                classifier=classifier,
                                param_col=param_col if 'param_col' in locals() else None,
                                value_col=value_col if 'value_col' in locals() else None,
                                purpose_filter=purpose_filter if 'purpose_filter' in locals() else None,
                                depth_filter=depth_filter if 'depth_filter' in locals() else None
                            )
                            
                            # Create download link
                            b64_pdf = base64.b64encode(pdf_content).decode()
                            href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="relatorio_solo_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf">{t("download_pdf")}</a>'
                            st.markdown(href, unsafe_allow_html=True)
                            
                            st.success(t('pdf_generated'))
                            
                        except Exception as e:
                            st.error(f"{t('pdf_error')} {str(e)}")
                
                with col2:
                    if st.button(t('generate_docx'), type="secondary"):
                        with st.spinner(t('generating_docx')):
                            try:
                                # Pass geospatial data and other parameters for kriging maps
                                docx_content = generate_docx_report(
                                    classified_df, 
                                    summary_stats, 
                                    None,  # charts_data placeholder
                                    st.session_state.get('project_name', t('default_project')),
                                    points_gdf=points_gdf if 'points_gdf' in locals() else None,
                                    polygon_gdf=polygon_gdf if 'polygon_gdf' in locals() else None,
                                    classifier=classifier,
                                    param_col=param_col if 'param_col' in locals() else None,
                                    value_col=value_col if 'value_col' in locals() else None,
                                    purpose_filter=purpose_filter if 'purpose_filter' in locals() else None,
                                    depth_filter=depth_filter if 'depth_filter' in locals() else None
                                )
                                
                                # Create download button for DOCX
                                st.download_button(
                                    label=t('download_docx'),
                                    data=docx_content,
                                    file_name=f"relatorio_solo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                
                                st.success(t('docx_generated'))
                                
                            except Exception as e:
                                st.error(f"{t('docx_error')} {str(e)}")
            
            with col3:
                    # Excel download with colors
                    output = io.BytesIO()
                    current_language = st.session_state.get('language', 'pt')
                    
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        # Prepare export dataframe
                        export_df = classified_df.copy()
                        export_df["Classificação"] = export_df["Classificação"].apply(translate_classification)
                        
                        # Translate to English if language is English
                        if current_language == 'en':
                            export_df = translate_column_names_to_english(export_df)
                            main_sheet_name = 'Classified_Data'
                            classification_col_name = 'Classification'
                        else:
                            main_sheet_name = 'Dados_Classificados'
                            classification_col_name = 'Classificação'
                        
                        export_df.to_excel(writer, sheet_name=main_sheet_name, index=False)
                        
                        # Apply colors to the classification column
                        apply_excel_colors(writer.book[main_sheet_name], classification_col_name)
                        
                        # Médias (Averages) sheet by plot_type
                        create_medias_sheet(writer, classified_df, param_col, value_col, language=current_language)
                        
                        # Statistics per plot_type, sampling_plan_purpose, depth_range_bottom_m with classified means
                        stats_df = create_comprehensive_statistics_with_classification(
                            classified_df, 
                            ['Tratamento', 'Data amostragem', 'profundidade inferior'], 
                            param_col, 
                            value_col,
                            language=current_language
                        )
                        
                        if stats_df is not None and not stats_df.empty:
                            if current_language == 'en':
                                stats_sheet_name = 'Detailed_Statistics'
                                stats_classification_col = 'Mean_Classification'
                            else:
                                stats_sheet_name = 'Estatísticas_Detalhadas'
                                stats_classification_col = 'Classificação_Média'
                                
                            stats_df.to_excel(writer, sheet_name=stats_sheet_name, index=False)
                            # Apply colors to the classification column in statistics
                            apply_excel_colors(writer.book[stats_sheet_name], stats_classification_col)
                        
                        # Create color legend sheet
                        create_color_legend_sheet(writer.book, language=current_language)
                    
                    excel_data = output.getvalue()
                    
                    st.download_button(
                        label=t('download_excel'),
                        data=excel_data,
                        file_name=f"dados_classificados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
        
        except Exception as e:
            st.error(f"❌ Erro ao processar arquivo: {str(e)}")
            st.info("Verifique se o arquivo está no formato correto e tente novamente.")
    
    else:
        # Instructions when no file is uploaded
        st.info(t('upload_instructions'))
        
        st.markdown(f"""
        ### {t('instructions')}
        
        1. {t('step1')}
        2. {t('step2')}
        3. {t('step3')}
        4. {t('step4')}
        5. {t('step5')}
        6. {t('step6')}
        
        ### {t('expected_format')}
        
        {t('format_description')}
        - {t('parameter_col_desc')}
        - {t('value_col_desc')}
        
        ### {t('supported_classifications')}
        
        - {t('mb_mbom')}
        - {t('b_malto')}
        """)

if __name__ == "__main__":
    main()

